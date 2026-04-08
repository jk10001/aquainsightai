# filename: report.py
from datetime import date
from typing import List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# LLM CHECKLIST (for context, not executed):
# 1. Configure A4 page layout, margins, header, and footer with date and page numbers.
# 2. Add title page and front matter (TOC, List of Figures, List of Tables).
# 3. Insert structured sections: Introduction, Objectives, Dataset Description, Analysis,
#    Key Findings, and Recommendations, with process-engineering-focused narrative.
# 4. Create compliance tables for flow, ammonia, and BOD using annual statistics vs limits.
# 5. Insert all provided charts as figures with automatic captions and reference them in text.
# 6. Summarise trends, anomalies (2017 ammonia), data quality issues, and propose further analysis.
# 7. Save the assembled document as report.docx in the current directory.


# ────────────────────────────────────────────────────────────────
#                       low-level helpers
# ────────────────────────────────────────────────────────────────
def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)  # A4 height in mm
        s.page_width = Mm(210)  # A4 width in mm
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(
    doc: Document, title: str, report_date: str | None = None
) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        # header
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        # footer
        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # date (left cell)
        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style

        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        # page number (right cell)
        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style

        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    title_para = doc.add_paragraph(title, style="Title")
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para = doc.add_paragraph(report_date, style="Subtitle")
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    if not rows:
        return
    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # repeat header row
    t.rows[0].repeat_header = True

    available_width_mm = 160
    num_cols = len(rows[0])
    if num_cols > 0:
        col_width_mm = available_width_mm / num_cols
        for col in t.columns:
            col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    doc.add_paragraph()


# ────────────────────────────────────────────────────────────────
#                       MAIN REPORT GENERATION
# ────────────────────────────────────────────────────────────────
def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Eastern Treatment Plant Discharge Compliance Assessment 2014–2018"
    today_str = date.today().strftime("%d %B %Y")

    # Cover & prelims
    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(
        doc,
        include_toc=True,
        include_lof=True,
        include_lot=True,
    )

    # ───────────────────────────────────────────────────────────
    # 1 Introduction
    # ───────────────────────────────────────────────────────────
    doc.add_heading("1 Introduction", level=1)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "This report presents a discharge limit compliance assessment for the Melbourne "
        "Water Eastern Treatment Plant (ETP) over the period 2014 to 2018. The assessment "
        "focuses on treated effluent discharged from the plant, comparing actual "
        "performance against the regulatory limits for flow, ammonia, and biochemical "
        "oxygen demand (BOD). Results are based on daily flow and effluent quality data "
        "and are summarised as annual statistics supported by detailed time series plots."
    )

    # ───────────────────────────────────────────────────────────
    # 2 Objectives
    # ───────────────────────────────────────────────────────────
    doc.add_heading("2 Objectives", level=1)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "The objectives of this assessment are to quantify the ETP’s compliance with "
        "licensed discharge limits and to identify any periods of non‑compliance or "
        "emerging risk that may warrant operational or process improvements."
    )

    doc.add_paragraph(
        "The specific objectives are:", style="Body Text"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph(
        "Tabulate annual mean flow and effluent quality statistics for ammonia and BOD "
        "against their respective discharge limits for 2014–2018.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Visualise annual compliance margins using bar charts for each regulated "
        "parameter and full‑period time series for flow, ammonia, and BOD.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Interpret trends and anomalies from a process engineering perspective, "
        "including likely causes of any excursions.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Provide key findings and recommendations for ongoing compliance management "
        "and further investigation.",
        style="List Bullet",
    )

    # ───────────────────────────────────────────────────────────
    # 3 Description of Dataset
    # ───────────────────────────────────────────────────────────
    doc.add_heading("3 Description of Dataset", level=1)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "The assessment uses five years of operational data from the Eastern Treatment "
        "Plant (1 January 2014 to 31 December 2018), extracted from the file "
        "MWC_ETP_Data.xlsx. Daily influent flow records and irregularly sampled "
        "effluent quality measurements for ammonia and BOD were analysed to derive "
        "annual statistics. Regulatory discharge limits were sourced from the "
        "“Discharge Limits” sheet and applied to the corresponding parameters. "
        "Electricity usage data are available in the dataset but are not used in this "
        "compliance assessment."
    )

    # ───────────────────────────────────────────────────────────
    # 4 Analysis
    # ───────────────────────────────────────────────────────────
    doc.add_heading("4 Analysis", level=1)

    # 4.1 Flow compliance
    doc.add_heading("4.1 Discharge Flow Compliance", level=2)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Annual mean discharge flow was calculated from the daily influent flow series "
        "and compared with the licensed annual mean limit of 540 ML/d. Flows ranged "
        "between approximately 337 and 358 ML/d, corresponding to 62–66% of the "
        "allowable limit. This indicates substantial available hydraulic capacity and "
        "comfortable compliance margins in all years, as summarised in Table 1."
    )

    # Flow compliance table (Table 1)
    flow_limit = 540.0
    flow_years = [2014, 2015, 2016, 2017, 2018]
    flow_means = [337.405, 342.521, 358.068, 354.452, 343.107]

    flow_rows: List[List[str]] = [
        [
            "Year",
            "Annual Mean Flow (ML/d)",
            "Discharge Limit (ML/d)",
            "Flow as % of Limit",
        ]
    ]
    for year, mean_val in zip(flow_years, flow_means):
        pct = 100.0 * mean_val / flow_limit
        flow_rows.append(
            [
                str(year),
                f"{mean_val:,.1f}",
                f"{flow_limit:,.0f}",
                f"{pct:,.1f}",
            ]
        )

    add_table(
        doc,
        "Annual mean discharge flow compared with the 540 ML/d licence limit.",
        flow_rows,
    )

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Daily flow behaviour is illustrated in Figure 1 and shows pronounced "
        "variability with recurrent wet‑weather peaks, including several events above "
        "800 ML/d and a major storm in January 2018 reaching approximately "
        "1,330 ML/d. Despite these short‑term surges, the annual means remain well "
        "below the limit, confirming that the plant’s hydraulic design and wet‑weather "
        "management strategies are effective."
    )

    # Figure 1 – Daily discharge flow time series
    try:
        doc.add_picture("daily_discharge_flow_2014_2018.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Daily influent flow to the Eastern Treatment Plant (2014–2018) with "
            "the 540 ML/d annual mean discharge limit indicated.",
        )
    except Exception:
        pass

    # 4.2 Ammonia compliance
    doc.add_heading("4.2 Effluent Ammonia Compliance", level=2)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Effluent ammonia is regulated by both an annual mean limit of 0.5 mg/L and a "
        "90th percentile limit of 2.0 mg/L. Annual statistics were computed from the "
        "irregularly sampled effluent ammonia dataset. The results are presented in "
        "Table 2."
    )

    # Ammonia compliance table (Table 2)
    amm_mean_limit = 0.5
    amm_p90_limit = 2.0
    amm_mean = [0.0907258, 0.0859339, 0.0938031, 1.32132, 0.127017]
    amm_p90 = [0.1315, 0.11, 0.13, 0.13, 0.18]

    ammonia_rows: List[List[str]] = [
        [
            "Year",
            "Annual Mean (mg/L)",
            "Mean as % of 0.5 mg/L Limit",
            "90th Percentile (mg/L)",
            "90th Percentile as % of 2.0 mg/L Limit",
        ]
    ]
    for year, mean_val, p90_val in zip(flow_years, amm_mean, amm_p90):
        mean_pct = 100.0 * mean_val / amm_mean_limit
        p90_pct = 100.0 * p90_val / amm_p90_limit
        ammonia_rows.append(
            [
                str(year),
                f"{mean_val:,.3f}",
                f"{mean_pct:,.1f}",
                f"{p90_val:,.3f}",
                f"{p90_pct:,.1f}",
            ]
        )

    add_table(
        doc,
        "Annual effluent ammonia statistics compared with the 0.5 mg/L mean and "
        "2.0 mg/L 90th percentile discharge limits.",
        ammonia_rows,
    )

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "For four of the five years (2014–2016 and 2018), the annual mean ammonia "
        "concentration remained between 0.086 and 0.128 mg/L, equivalent to only "
        "17–25% of the 0.5 mg/L limit. In all years, the 90th percentile values "
        "remained very low at 0.11–0.18 mg/L (6–9% of the 2.0 mg/L limit). These "
        "results confirm consistently strong nitrification performance except in 2017, "
        "when the annual mean increased to 1.32 mg/L (264% of the limit) while the "
        "90th percentile remained at 0.13 mg/L. This combination points to a limited "
        "number of very high ammonia events rather than a sustained loss of process "
        "performance."
    )

    # Figure 2 – Annual ammonia statistics
    try:
        doc.add_picture(
            "annual_ammonia_mean_90th_percentile_vs_limits.png", width=Mm(160)
        )
        add_caption(
            doc,
            "Figure",
            "Annual mean and 90th percentile effluent ammonia concentrations "
            "compared with licence limits (2014–2018).",
        )
    except Exception:
        pass

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "The underlying daily effluent ammonia time series in Figure 3 shows that "
        "ammonia was typically below 0.5 mg/L, but several major nitrification "
        "failures occurred. A catastrophic event in mid‑2017 produced a single "
        "measurement around 316 mg/L, accompanied by multiple excursions between "
        "3 and 10 mg/L in 2017 and 2018. Isolated spikes also appear early in 2014. "
        "These episodic failures explain the elevated 2017 annual mean while the "
        "90th percentile remained low."
    )

    # Figure 3 – Daily effluent ammonia
    try:
        doc.add_picture(
            "daily_effluent_ammonia_2014_2018_v2.png",
            width=Mm(160),
        )
        add_caption(
            doc,
            "Figure",
            "Daily effluent ammonia concentration with annual mean and "
            "90th percentile limits (2014–2018).",
        )
    except Exception:
        pass

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "From a process engineering perspective, the pattern suggests intermittent "
        "nitrification upsets rather than chronic under‑design. Potential causes "
        "include low dissolved oxygen in aeration tanks during peak load events, "
        "temperature‑related nitrifier inhibition in winter, sudden toxic shocks from "
        "industrial discharges, or insufficient active biomass following sludge "
        "wasting or maintenance activities. The recurrence of elevated events in "
        "2017–2018 indicates that additional process control measures may be required."
    )

    # 4.3 BOD compliance
    doc.add_heading("4.3 Effluent BOD Compliance", level=2)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Effluent BOD is controlled by a 90th percentile limit of 10 mg/L. Annual "
        "statistics were derived from the effluent BOD measurements. The results are "
        "summarised in Table 3."
    )

    # BOD compliance table (Table 3)
    bod_p90_limit = 10.0
    bod_p90 = [3.0, 4.0, 4.0, 3.0, 3.0]

    bod_rows: List[List[str]] = [
        [
            "Year",
            "BOD 90th Percentile (mg/L)",
            "90th Percentile as % of 10 mg/L Limit",
            "Compliance Comment",
        ]
    ]
    for year, p90_val in zip(flow_years, bod_p90):
        pct = 100.0 * p90_val / bod_p90_limit
        comment = "Within limit with large margin"
        bod_rows.append(
            [
                str(year),
                f"{p90_val:,.1f}",
                f"{pct:,.1f}",
                comment,
            ]
        )

    add_table(
        doc,
        "Annual effluent BOD 90th percentile compared with the 10 mg/L discharge limit.",
        bod_rows,
    )

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Across all five years, the BOD 90th percentile remained between 3 and 4 mg/L, "
        "corresponding to only 30–40% of the 10 mg/L limit. This demonstrates very "
        "robust carbonaceous removal and process stability in the secondary "
        "treatment system."
    )

    # Figure 4 – Annual BOD 90th percentile
    try:
        doc.add_picture(
            "annual_bod_90th_percentile_vs_limit.png",
            width=Mm(160),
        )
        add_caption(
            doc,
            "Figure",
            "Annual effluent BOD 90th percentile compared with the 10 mg/L "
            "discharge limit (2014–2018).",
        )
    except Exception:
        pass

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Figure 5 presents the full effluent BOD time series. Measurements cluster at "
        "or near the typical detection limit of 2 mg/L, with occasional excursions up "
        "to 5–6 mg/L distributed across the monitoring period. No measurements "
        "approach the 10 mg/L limit, and there is no evidence of sustained process "
        "deterioration, confirming excellent and consistent organic removal."
    )

    # Figure 5 – Daily effluent BOD
    try:
        doc.add_picture(
            "daily_effluent_bod_2014_2018.png",
            width=Mm(160),
        )
        add_caption(
            doc,
            "Figure",
            "Daily effluent BOD concentrations with the 10 mg/L 90th percentile "
            "limit (2014–2018).",
        )
    except Exception:
        pass

    # 4.4 Integrated performance assessment
    doc.add_heading("4.4 Integrated Performance and Process Interpretation", level=2)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "When considered together, the flow, ammonia, and BOD results indicate that "
        "the Eastern Treatment Plant generally operates with substantial capacity and "
        "strong process robustness. Hydraulic loads are well below the design limit, "
        "and BOD performance is consistently excellent. The only material compliance "
        "concern identified is the 2017 ammonia annual mean exceedance, driven by a "
        "small number of extreme events."
    )

    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "The lack of corresponding deterioration in BOD performance during the high "
        "ammonia episodes suggests that the organic removal pathway remained intact "
        "while the specialised nitrifying biomass was temporarily impaired. This is "
        "consistent with short‑term inhibition of nitrifiers due to low dissolved "
        "oxygen, toxic compounds, or rapid changes in loading or temperature. Given "
        "the recurrence of elevated ammonia events in 2018, it is likely that "
        "underlying causes were not fully resolved following the 2017 incidents."
    )

    # 4.5 Data quality and anomalies
    doc.add_heading("4.5 Data Quality Considerations", level=2)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "The daily flow series contains a small number of anomalous negative or very "
        "low values in early 2014, which are not physically plausible and most likely "
        "reflect data logging or transcription errors. These points have negligible "
        "impact on annual statistics but should be flagged for correction in the "
        "source systems. Similarly, the extremely high ammonia concentration in "
        "mid‑2017 should be verified as a genuine event (for example, by cross‑checking "
        "laboratory records and concurrent process data) before being used in any "
        "design or risk assessments."
    )

    # 4.6 Opportunities for further analysis
    doc.add_heading("4.6 Further Analysis Opportunities", level=2)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Several additional analyses could be performed to deepen understanding of "
        "the plant’s performance and the root causes of the ammonia excursions."
    )

    doc.add_paragraph(
        "Undertake seasonal analysis of ammonia and BOD performance to separate "
        "temperature effects from load‑related impacts.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Correlate ammonia spikes with influent load, flow peaks, and electricity "
        "usage (as a proxy for aeration intensity) to identify potential oxygen "
        "limitation or shock loading.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Review process events during 2017–2018 (maintenance, bypasses, industrial "
        "trade waste incidents) using plant logs and supervisory control and data "
        "acquisition (SCADA) records.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Evaluate nitrification safety margins under design‑storm conditions using "
        "dynamic process modelling calibrated to the observed data.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Assess redundancy and resilience of aeration and mixed liquor recycle "
        "systems to confirm the plant’s ability to manage extreme events without "
        "breaching ammonia limits.",
        style="List Bullet",
    )

    # ───────────────────────────────────────────────────────────
    # 5 Key Findings
    # ───────────────────────────────────────────────────────────
    doc.add_heading("5 Key Findings", level=1)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "The main findings from the 2014–2018 discharge compliance assessment are "
        "summarised below."
    )

    doc.add_paragraph(
        "Annual mean flows were consistently between 337 and 358 ML/d, representing "
        "62–66% of the 540 ML/d licence limit, with significant remaining hydraulic "
        "capacity.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Effluent BOD performance was excellent, with 90th percentile values of "
        "3–4 mg/L (30–40% of the 10 mg/L limit) and daily data largely at the "
        "analytical detection limit, indicating robust biological treatment.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Effluent ammonia performance was generally very good, with annual mean "
        "concentrations below 0.13 mg/L and 90th percentile values below 0.18 mg/L "
        "for four of the five years.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "In 2017 the annual mean ammonia concentration increased to 1.32 mg/L "
        "(264% of the limit) due to a small number of extreme nitrification failures, "
        "while the 90th percentile remained compliant.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "The ammonia upsets do not coincide with any observable deterioration in BOD "
        "performance, suggesting selective impairment of nitrifying biomass rather "
        "than global process failure.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Data quality is generally good, but anomalous negative flow readings and "
        "extreme ammonia outliers should be validated and, where appropriate, "
        "corrected or annotated.",
        style="List Bullet",
    )

    # ───────────────────────────────────────────────────────────
    # 6 Recommendations
    # ───────────────────────────────────────────────────────────
    doc.add_heading("6 Recommendations", level=1)
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
        "Based on the analysis and findings, the following recommendations are "
        "proposed to maintain and enhance discharge compliance at the Eastern "
        "Treatment Plant."
    )

    doc.add_paragraph(
        "Investigate the 2017–2018 ammonia excursions in detail, including review of "
        "SCADA trends, plant logs, and trade waste records, to confirm root causes "
        "such as dissolved oxygen limitation, toxic shocks, or operational "
        "constraints.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Develop and implement targeted nitrification resilience measures, such as "
        "improved aeration control strategies, minimum dissolved oxygen set‑points, "
        "and contingency protocols for high‑risk industrial discharges.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Formalise data validation procedures for flow and effluent quality data to "
        "systematically identify and correct non‑physical values before compliance "
        "reporting.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Continue routine performance monitoring using annual statistics and time "
        "series plots similar to those in this report to detect emerging trends "
        "early.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Consider complementary monitoring of key process parameters (for example, "
        "dissolved oxygen, mixed liquor suspended solids, and sludge age) to support "
        "proactive management of nitrification capacity.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Use dynamic process modelling or targeted pilot trials to assess the plant’s "
        "response to extreme wet‑weather and high‑load events and to optimise "
        "operational strategies for maintaining ammonia compliance.",
        style="List Bullet",
    )

    # FINAL STEP: Save the document to disk.
    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()