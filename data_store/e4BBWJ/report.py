# filename: report.py
"""
Generates a professional DOCX report for:
Ontario Surface Water Quality Screening Assessment - 2024

Revision updates implemented per user instructions:
- Removed the preparation checklist section from the report.
- Table 1 caption changed to: “Analyte summary (PWQO screening, 2024)”.
- Table 5 rebuilt to include separate columns for:
  Analyte, Objective, N Samples, N exceedances, % exceedances,
  N stations >=1 exceedance, % non-detect, median, and 95th percentile.
  Table 5 uses 9pt font and caption: “Detailed analytes screening summary”.

Notes:
- Uses python-docx only (no pandas).
- Inserts updatable fields for TOC, List of Figures, List of Tables.
- Adds automatically numbered captions for all figures and tables.
- Uses A4 paper; Appendix A is landscape for readability.
"""

from __future__ import annotations

from datetime import date
from typing import List, Optional
import os
import math

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# ────────────────────────────────────────────────────────────────
# Low-level helpers
# ────────────────────────────────────────────────────────────────
def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)  # A4
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Optional[Pt] = None,
    font_color: Optional[RGBColor] = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)
    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(doc: Document, title: str, report_date: Optional[str] = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    # Ensure footer style is 10pt grey so field results match.
    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        # Header (center)
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        # Footer (2-column table: date left, page number right)
        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        right_para = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            right_para.style = footer_style
        right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(right_para, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: Optional[str] = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    t = doc.add_paragraph(title, style="Title")
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    d = doc.add_paragraph(report_date, style="Subtitle")
    d.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        p = doc.add_paragraph()
        insert_field(p, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        p = doc.add_paragraph()
        insert_field(p, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        p = doc.add_paragraph()
        insert_field(p, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def _set_table_repeat_header_row(table) -> None:
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
    font_size_pt: int = 10,
    max_width_mm: float = 160.0,
) -> None:
    if not rows or not rows[0]:
        return

    add_caption(doc, "Table", caption)

    n_rows = len(rows)
    n_cols = len(rows[0])
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_table_repeat_header_row(t)

    col_w = max_width_mm / max(n_cols, 1)
    for col in t.columns:
        col.width = Mm(col_w)

    for r in range(n_rows):
        if len(rows[r]) != n_cols:
            raise ValueError("All table rows must have the same number of columns.")
        for c in range(n_cols):
            t.cell(r, c).text = str(rows[r][c])

    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size_pt)

    doc.add_paragraph()  # spacer after table


def add_figure(
    doc: Document,
    image_path: str,
    caption: str,
    width_mm: float = 160.0,
) -> None:
    if not os.path.exists(image_path):
        p = doc.add_paragraph(
            f"Figure not available in working directory: {image_path}.",
            style="Body Text",
        )
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        add_caption(doc, "Figure", caption)
        doc.add_paragraph()
        return

    doc.add_picture(image_path, width=Mm(width_mm))
    add_caption(doc, "Figure", caption)
    doc.add_paragraph()


def add_landscape_section(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Mm(297)
    sec.page_height = Mm(210)
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(15)
    sec.right_margin = Mm(15)


def add_paragraph_justify(doc: Document, text: str, style: str = "Body Text") -> Paragraph:
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p


def fmt_num(x: float, decimals: int = 1) -> str:
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return ""
    return f"{x:,.{decimals}f}"


# ────────────────────────────────────────────────────────────────
# Report content
# ────────────────────────────────────────────────────────────────
def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Ontario Surface Water Quality Screening Assessment - 2024"
    today_str = date.today().strftime("%d %B %Y")

    # Cover
    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)

    # Front matter
    add_front_matter(doc, include_toc=True, include_lof=True, include_lot=True)

    # -----------------------------
    # 1. Introduction
    # -----------------------------
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Background", level=2)
    add_paragraph_justify(
        doc,
        "The Provincial (Stream) Water Quality Monitoring Network (PWQMN) measures surface-water quality in rivers and streams across Ontario. "
        "This report summarises a 2024 screening assessment of PWQMN results against applicable Provincial Water Quality Objectives (PWQO) "
        "and Interim PWQO values for analytes with province-wide objectives.",
    )

    doc.add_heading("1.2 Scope of Assessment", level=2)
    add_paragraph_justify(
        doc,
        "This assessment is a screening-level evaluation based on individual sample results collected between 05 January 2024 and 11 December 2024. "
        "The approach flags exceedance signals on a per-sample basis and then aggregates results across stations and analytes to support prioritisation. "
        "It does not replace the formal application requirements associated with some objectives (e.g., multi-sample geometric means).",
    )

    doc.add_heading("1.3 Focus Analytes (Phosphorus and Iron)", level=2)
    add_paragraph_justify(
        doc,
        "While multiple analytes are screened where PWQO (or Interim PWQO) values are available, this report places additional focus on Total Phosphorus "
        "and Iron because they are among the most frequently exceeded objectives and are strongly linked to common watershed and process drivers "
        "(e.g., erosion/sediment mobilisation, runoff, and nutrient loading).",
    )

    doc.add_heading("1.4 Regulatory Context (Ontario PWQO Framework)", level=2)
    add_paragraph_justify(
        doc,
        "PWQO values provide environmental quality benchmarks intended to protect aquatic life and other water uses under typical conditions. "
        "Objective applicability depends on parameter-specific requirements (e.g., pH-dependent objectives, hardness-adjusted metals, or geometric mean criteria). "
        "Accordingly, results in this report are framed as a conservative screen to identify potential areas of concern requiring follow-up.",
    )

    # -----------------------------
    # 2. Data and Methods
    # -----------------------------
    doc.add_heading("2. Data and Methods", level=1)

    doc.add_heading("2.1 Dataset Description", level=2)
    add_paragraph_justify(
        doc,
        "The source dataset is an Excel workbook (Ontario_PWQMN_2024.xlsx) with two worksheets: "
        "(i) a Data sheet containing individual sample results by station and analyte, including result qualifiers and detection limits, and "
        "(ii) a Stations sheet containing station metadata (name, latitude/longitude, and status).",
    )

    doc.add_heading("2.2 Treatment of Non-Detects", level=2)
    add_paragraph_justify(
        doc,
        "Non-detect (below detection limit) results were handled using a simple substitution approach for summary statistics: non-detects were assigned "
        "a value of one-half the method detection limit (½MDL) when calculating mean, median, and percentile metrics. "
        "For exceedance screening, exceedances were evaluated on detected results as a conservative indicator of objective exceedance signals.",
    )

    doc.add_heading("2.3 PWQO Objectives Applied", level=2)
    add_paragraph_justify(
        doc,
        "Screening was performed only for analytes with a provincial objective (PWQO or Interim PWQO). "
        "Where objectives are conditional (e.g., hardness-adjusted metals or pH-dependent objectives), screening required the relevant supporting data "
        "(e.g., same-day hardness or pH). Where supporting data were missing, those samples were excluded from exceedance screening for that parameter.",
    )

    doc.add_heading("2.4 Exceedance Definition", level=2)
    add_paragraph_justify(
        doc,
        "A sample exceedance was defined as a measured (detected) result above the applicable objective threshold. "
        "For pH, an exceedance was defined as values outside the PWQO range of 6.5 to 8.5. "
        "For E. coli, samples above 100 MPN/100 mL were flagged; this is more conservative than the formal PWQO application, which is based on a geometric mean of at least five samples.",
    )

    doc.add_heading("2.5 Limitations of Screening Approach", level=2)
    add_paragraph_justify(
        doc,
        "Key limitations relevant to interpreting screening results include: per-sample screening does not incorporate flow conditions or event context; "
        "some objectives have formal aggregation requirements not applied here; and detection limits can prevent meaningful assessment when objectives are lower than MDLs. "
        "Consequently, the output should be used to prioritise follow-up investigations rather than to conclude definitive compliance status.",
    )

    add_paragraph_justify(
        doc,
        "The next section summarises screening results across the network, followed by a risk-tier prioritisation framework and detailed analyses for Total Phosphorus and Iron.",
    )

    # -----------------------------
    # 3. Provincial Summary Results
    # -----------------------------
    doc.add_heading("3. Provincial Summary Results", level=1)

    add_paragraph_justify(
        doc,
        "Network-wide screening results indicate that objective exceedances are concentrated in a subset of analytes and stations. "
        "Table 1 summarises results for all PWQO-screened analytes, including objectives, sample counts, exceedance rates (detected-only screen), "
        "and upper-tail concentrations (P95). Detailed analyte screening and censoring metrics are provided in Appendix A (Table 5).",
    )

    table1 = [
        ["Analyte", "Objective", "N samples", "Exceedances [%]", "P95"],
        ["Aluminum", "pH-dependent", "2,134", "42.5", "643.0 µg/L"],
        ["Arsenic", "5.0 µg/L", "199", "14.1", "21.1 µg/L"],
        ["Boron", "200.0 µg/L", "93", "0.0", "75.6 µg/L"],
        ["Cadmium", "Hardness-adjusted", "2,134", "5.2", "1.01 µg/L"],
        ["Chromium", "1.0 µg/L", "2,070", "9.7", "2.5 µg/L"],
        ["Copper", "Hardness-adjusted", "2,133", "9.3", "6.552 µg/L"],
        ["E. coli", "100 MPN/100 mL", "216", "56.9", "3,275.0"],
        ["Iron", "300.0 µg/L", "2,132", "28.8", "919.8 µg/L"],
        ["Lead", "Hardness-adjusted", "2,132", "0.0", "3.5 µg/L"],
        ["Mercury", "0.2 µg/L", "142", "0.0", "2.5 ng/L (½MDL)"],
        ["Nickel", "25.0 µg/L", "2,134", "1.5", "6.277 µg/L"],
        ["Total Phosphorus", "30.0 µg/L", "3,666", "40.1", "237.75 µg/L"],
        ["Selenium", "100.0 µg/L", "95", "0.0", "2.5 µg/L (½MDL)"],
        ["Silver", "0.1 µg/L", "2,117", "0.0", "4.5 µg/L (½MDL)"],
        ["Zinc", "20.0 µg/L", "2,114", "8.6", "29.835 µg/L"],
        ["pH", "6.5–8.5 SU", "3,669", "4.3", "8.49"],
    ]
    add_table(
        doc,
        "Analyte summary (PWQO screening, 2024)",
        table1,
        font_size_pt=9,
        max_width_mm=160.0,
    )

    add_paragraph_justify(
        doc,
        "At the provincial scale, the strongest per-sample exceedance signals were observed for E. coli (56.9% of 216 samples), Aluminum (42.5% of 2,134 samples), "
        "Total Phosphorus (40.1% of 3,666 samples), and Iron (28.8% of 2,132 samples). For Total Phosphorus and Iron, the combination of high exceedance frequency and "
        "elevated upper-tail concentrations indicates that exceedances are not restricted to rare, extreme outliers and warrants station-level prioritisation.",
    )

    add_paragraph_justify(
        doc,
        "Detection-limit constraints can materially affect the ability to assess compliance. A notable example is Silver: 100% of results were non-detect and 100% of method "
        "detection limits exceeded the 0.1 µg/L objective, meaning that compliance cannot be concluded from these data. Where the objective is below typical MDLs, future monitoring "
        "may require lower reporting limits or targeted analytical methods to support meaningful assessment.",
    )

    # -----------------------------
    # 4. Prioritisation Risk Framework and Results
    # -----------------------------
    doc.add_heading("4. Prioritisation Risk Framework and Results", level=1)

    doc.add_heading("4.1 Risk Classification Framework", level=2)
    add_paragraph_justify(
        doc,
        "Stations were classified into screening risk tiers based on the overall percent of PWQO-screened sample results exceeding objectives across all objective-screened analytes. "
        "The tier definitions are summarised in Table 2.",
    )

    doc.add_heading("4.2 Risk Classification Results", level=2)
    add_paragraph_justify(
        doc,
        "Figure 1 shows the geographic distribution of station risk tiers. Tier 3–4 stations are concentrated in Southern Ontario, consistent with higher pressures from intensive land use, "
        "urban runoff, and agricultural drainage in several watersheds. Table 2 summarises the number of stations in each tier for the screened station set (n=435), together with the "
        "exceedance criteria used to define each tier.",
    )

    add_figure(
        doc,
        "pwqmn_stations_map_risk_2024.png",
        "PWQMN 2024 stations map (risk tier).",
        width_mm=160.0,
    )

    tier_counts_table = [
        ["Risk tier", "Exceedance criterion", "Stations", "Share [%]"],
        ["Tier 1", "0% of screened results exceed objectives", "47", "10.8"],
        ["Tier 2", "≤15% of screened results exceed objectives", "186", "42.8"],
        ["Tier 3", "15–50% of screened results exceed objectives", "191", "43.9"],
        ["Tier 4", "≥50% of screened results exceed objectives", "11", "2.5"],
        ["Total", "All tiers", "435", "100.0"],
    ]
    add_table(
        doc,
        "Station counts by risk tier and screening exceedance criteria (PWQO screening, 2024).",
        tier_counts_table,
        font_size_pt=9,
        max_width_mm=160.0,
    )

    add_paragraph_justify(
        doc,
        "The station-level results used for tiering include parameter-by-parameter exceedance summaries. The full station table is provided as a CSV file: "
        "station_risk_classification_pwqo_2024_v2.csv.",
    )

    # -----------------------------
    # 5. Phosphorus Analysis
    # -----------------------------
    doc.add_heading("5. Phosphorus Analysis", level=1)
    add_paragraph_justify(
        doc,
        "Total Phosphorus was screened against the river PWQO of 30 µg/L. At the provincial scale, 40.1% of screened samples exceeded the objective "
        "(1,469 exceedances of 3,666 samples), and 312 stations recorded at least one exceedance during 2024. Figure 2 maps station-level exceedance "
        "frequency, highlighting clusters of persistent exceedance in Southern Ontario.",
    )

    add_figure(
        doc,
        "tp_exceedance_map_2024.png",
        "PWQMN 2024 stations map (Total Phosphorus percent exceedance of PWQO).",
        width_mm=160.0,
    )

    add_paragraph_justify(
        doc,
        "Table 3 lists the top 10 stations by Total Phosphorus exceedance performance. All listed stations show 100% per-sample exceedance, indicating "
        "a persistent screening signal rather than a small number of isolated event-driven samples. These stations are efficient candidates for follow-up "
        "investigations, including upstream source screening (point sources, tile drainage, and stormwater inputs) and pairing Total Phosphorus with turbidity "
        "or suspended solids to distinguish particulate-bound versus dissolved contributions.",
    )

    tp_top10 = [
        ["Station", "N", "Median [µg/L]", "P95 [µg/L]", "Max [µg/L]"],
        ["16002700102 (Sturgeon River)", "10", "1,135.0", "4,494.5", "5,030.0"],
        ["16003000302 (Lebo Drain)", "7", "737.0", "963.8", "974.0"],
        ["11000101302 (Oswego Creek)", "10", "562.5", "1,051.1", "1,290.0"],
        ["4001303302 (Big Creek)", "6", "466.0", "2,162.5", "2,400.0"],
        ["6000300102 (Four Mile Creek)", "9", "364.0", "3,500.8", "5,400.0"],
        ["4001000302 (Ruscom River)", "10", "331.0", "1,947.4", "2,860.0"],
        ["11000100702 (Welland River)", "10", "278.0", "496.9", "500.0"],
        ["6002400802 (Twenty Mile Creek)", "14", "242.5", "590.2", "598.0"],
        ["4001308102 (Mcgregor Creek)", "7", "236.0", "507.1", "565.0"],
        ["16007200102 (Dutton Drain)", "7", "216.0", "291.8", "293.0"],
    ]
    add_table(doc, "Top 10 stations – Total Phosphorus PWQO exceedance (2024).", tp_top10)

    add_paragraph_justify(
        doc,
        "Figure 3 summarises monthly Total Phosphorus distributions (all sites combined). Medians are near the 30 µg/L PWQO in most months, while the upper whiskers "
        "exceed the PWQO throughout the year. This indicates that exceedances are frequent across the network and not solely driven by extreme outliers. Seasonal structure is visible "
        "in the spread: winter and mid-summer months show broader interquartile ranges, consistent with a combination of event-driven runoff (including snowmelt and storm events) and "
        "summer conditions that can amplify concentration variability through lower dilution and biological cycling.",
    )

    add_figure(
        doc,
        "monthly_boxplot_total_phosphorus_2024_v3.png",
        "Monthly box plot – Total Phosphorus (all sites combined, 2024).",
        width_mm=160.0,
    )

    # -----------------------------
    # 6. Iron Analysis
    # -----------------------------
    doc.add_heading("6. Iron Analysis", level=1)
    add_paragraph_justify(
        doc,
        "Iron was screened against the PWQO of 300 µg/L. Provincially, 28.8% of screened samples exceeded the objective (613 exceedances of 2,132 samples), and "
        "186 stations recorded at least one exceedance. Figure 4 maps station-level exceedance frequency, with higher frequencies concentrated in Southern Ontario.",
    )

    add_figure(
        doc,
        "iron_exceedance_map_2024.png",
        "PWQMN 2024 stations map (Iron percent exceedance of PWQO).",
        width_mm=160.0,
    )

    add_paragraph_justify(
        doc,
        "Table 4 lists the top 10 stations by iron exceedance performance. All listed stations show 100% per-sample exceedance, indicating persistent elevation at those sites. "
        "In practice, elevated iron can reflect particulate mobilisation (erosion, resuspension, and fine sediment transport) and/or natural sources such as groundwater influence "
        "in iron-rich geologies or peat-influenced catchments. Differentiating dissolved versus particulate iron is essential for selecting effective management responses.",
    )

    iron_top10 = [
        ["Station", "N", "Median [µg/L]", "P95 [µg/L]", "Max [µg/L]"],
        ["18207000802 (Cobbs Lake Creek)", "5", "1,570.0", "2,526.0", "2,660.0"],
        ["3013400302 (Emery Creek)", "7", "1,320.0", "2,385.0", "2,400.0"],
        ["1010700402 (Neebing River)", "8", "1,085.0", "1,899.5", "1,910.0"],
        ["13000300102 (Big Carp River)", "6", "1,045.0", "2,132.5", "2,450.0"],
        ["3013301502 (La Vase River)", "7", "976.0", "1,519.0", "1,630.0"],
        ["11000101302 (Oswego Creek)", "7", "967.0", "1,849.0", "1,960.0"],
        ["11000100702 (Welland River)", "7", "948.0", "1,480.0", "1,540.0"],
        ["3013301302 (Duchesnay River)", "7", "803.0", "1,071.0", "1,080.0"],
        ["16018409302 (Fairchild Creek)", "6", "756.5", "971.2", "1,020.0"],
        ["14002802802 (Whitson River)", "11", "667.0", "1,118.0", "1,280.0"],
    ]
    add_table(doc, "Top 10 stations – Iron PWQO exceedance (2024).", iron_top10)

    add_paragraph_justify(
        doc,
        "Figure 5 summarises monthly iron distributions (all sites combined). Medians are typically below 300 µg/L, indicating that typical conditions at many sites are compliant; "
        "however, the upper quartile frequently approaches or exceeds the PWQO, and upper whiskers exceed the PWQO in every month. The broader spreads in July and December are "
        "consistent with a combination of event/flow-driven mobilisation and seasonal changes in water chemistry and dilution. Pairing iron results with turbidity or suspended solids "
        "and sample flow condition (baseflow versus stormflow) would improve attribution of exceedance drivers.",
    )

    add_figure(
        doc,
        "monthly_boxplot_iron_2024.png",
        "Monthly box plot – Iron (all sites combined, 2024).",
        width_mm=160.0,
    )

    add_paragraph_justify(
        doc,
        "Figure 6 provides an all-sample time series view (log scale). The PWQO line is crossed throughout the year, supporting the interpretation that iron exceedances are not confined "
        "to one season and are driven by a mixture of baseline station characteristics and episodic events.",
    )

    add_figure(
        doc,
        "iron_timeseries_2024.png",
        "Time series – Iron (all samples, 2024; log scale).",
        width_mm=160.0,
    )

    # -----------------------------
    # 7. Conclusions and Recommendation
    # -----------------------------
    doc.add_heading("7. Conclusions and Recommendation", level=1)

    add_paragraph_justify(
        doc,
        "This 2024 screening assessment identified widespread objective exceedance signals for Total Phosphorus and Iron across the PWQMN network, with clear spatial clustering "
        "of higher-risk stations in Southern Ontario. The risk-tier framework (Table 2 and Figure 1) supports efficient prioritisation by focusing follow-up on Tier 4 and upper Tier 3 sites "
        "where chronic per-sample exceedances are most likely.",
    )

    add_paragraph_justify(doc, "Key conclusions from the 2024 screen include:")
    doc.add_paragraph(
        "Total Phosphorus shows a strong, network-wide exceedance signal (40.1% per-sample exceedance) and multiple stations with 100% exceedance and high medians (Table 3), indicating persistent enrichment at priority sites.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Iron exceedances (28.8% per-sample exceedance) are widespread, with several stations exhibiting 100% exceedance and high upper tails (Table 4), consistent with both sediment mobilisation and possible natural-source contributions.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Detection limits can prevent compliance assessment for some analytes (notably Silver), requiring improved reporting limits if compliance screening is desired for those parameters.",
        style="List Bullet",
    )
    doc.add_paragraph()

    add_paragraph_justify(doc, "Recommended follow-up actions (highest value first) are:")
    doc.add_paragraph(
        "Tier 4 and upper Tier 3 station review: confirm repeatability across years, review upstream land use and potential point sources, and confirm whether sampling is biased toward storm events or specific seasons.",
        style="List Number",
    )
    doc.add_paragraph(
        "Phosphorus source attribution at priority sites: pair Total Phosphorus with suspended solids/turbidity and, where feasible, dissolved (filtered) phosphorus fractions to determine whether exceedances are predominantly particulate-driven.",
        style="List Number",
    )
    doc.add_paragraph(
        "Iron speciation and transport assessment: add dissolved/filtered iron where feasible and routinely capture turbidity or suspended solids at sampling to distinguish sediment-associated transport from groundwater-driven dissolved iron.",
        style="List Number",
    )
    doc.add_paragraph(
        "Objective-application refinements: apply formal objective requirements where applicable (e.g., geometric mean criteria for E. coli where sufficient sample counts exist), and evaluate sensitivity to non-detect treatment choices for summary statistics.",
        style="List Number",
    )
    doc.add_paragraph(
        "Address detection-limit constraints: for analytes where objectives fall below typical MDLs, evaluate laboratory methods and reporting limits needed to support meaningful compliance assessment.",
        style="List Number",
    )
    doc.add_paragraph()

    add_paragraph_justify(
        doc,
        "Additional analyses that could further strengthen interpretation include trend comparisons to prior years, watershed-scale aggregation by basin, event versus baseflow stratification (where flow data are available), "
        "and multivariate screening (e.g., correlation of Total Phosphorus and Iron exceedances with suspended solids and turbidity).",
    )

    # -----------------------------
    # Appendix A (Landscape)
    # -----------------------------
    doc.add_heading("Appendix A. Detailed Analytes Table", level=1)
    add_paragraph_justify(
        doc,
        "Appendix A provides detailed screening metrics by analyte for parameters with provincial objectives. Statistics are based on the ½MDL substitution approach for non-detect results as described in Section 2.2.",
    )

    add_landscape_section(doc)
    doc.add_heading("Appendix A. Detailed Analytes Table", level=1)

    add_paragraph_justify(
        doc,
        "Table 5 summarises detailed screening metrics by analyte. For pH-dependent and hardness-adjusted objectives, screening was applied only where required supporting data were available.",
    )

    # Table 5: Separate columns (per user instruction); 9pt font
    table5 = [
        [
            "Analyte",
            "Objective",
            "N Samples",
            "N exceedances",
            "% exceedances",
            "N stations >=1 exceedance",
            "% non-detect",
            "Median",
            "95th percentile",
        ],
        ["Aluminum", "pH-dependent", "2,134", "906", "42.5", "218", "3.0", "56.7 µg/L", "643.0 µg/L"],
        ["Arsenic", "5.0 µg/L", "199", "28", "14.1", "7", "39.2", "1.0 µg/L", "21.1 µg/L"],
        ["Boron", "200.0 µg/L", "93", "0", "0.0", "0", "48.4", "10.0 µg/L", "75.6 µg/L"],
        ["Cadmium", "Hardness-adjusted", "2,134", "112", "5.2", "86", "94.8", "0.45 µg/L", "1.01 µg/L"],
        ["Chromium", "1.0 µg/L", "2,070", "200", "9.7", "113", "90.0", "0.5 µg/L", "2.5 µg/L"],
        ["Copper", "Hardness-adjusted", "2,133", "199", "9.3", "80", "8.0", "1.79 µg/L", "6.552 µg/L"],
        ["E. coli", "100 MPN/100 mL", "216", "123", "56.9", "21", "0.0", "131.5", "3,275.0"],
        ["Iron", "300.0 µg/L", "2,132", "613", "28.8", "186", "0.4", "168.0 µg/L", "919.8 µg/L"],
        ["Lead", "Hardness-adjusted", "2,132", "1", "0.0", "1", "99.3", "3.5 µg/L", "3.5 µg/L"],
        ["Mercury", "0.2 µg/L", "142", "0", "0.0", "0", "95.8", "2.5 ng/L", "2.5 ng/L"],
        ["Nickel", "25.0 µg/L", "2,134", "31", "1.5", "7", "82.1", "1.0 µg/L", "6.277 µg/L"],
        ["Total Phosphorus", "30.0 µg/L", "3,666", "1,469", "40.1", "312", "0.1", "23.6 µg/L", "237.75 µg/L"],
        ["Selenium", "100.0 µg/L", "95", "0", "0.0", "0", "100.0", "2.5 µg/L", "2.5 µg/L"],
        ["Silver", "0.1 µg/L", "2,117", "0", "0.0", "0", "100.0", "4.5 µg/L", "4.5 µg/L"],
        ["Zinc", "20.0 µg/L", "2,114", "182", "8.6", "76", "5.8", "4.86 µg/L", "29.835 µg/L"],
        ["pH", "6.5–8.5 SU", "3,669", "157", "4.3", "90", "0.0", "8.23", "8.49"],
    ]

    add_table(
        doc,
        "Detailed analytes screening summary",
        table5,
        font_size_pt=9,
        max_width_mm=260.0,
    )

    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()