# filename: report.py
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Inches
from docx.text.paragraph import Paragraph
from typing import List


def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(
    doc: Document, title: str, report_date: str | None = None
) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style
        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    title_para = doc.add_paragraph(title, style="Title")
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para = doc.add_paragraph(report_date, style="Subtitle")
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    if not rows:
        return
    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    t.rows[0].repeat_header = True

    available_width_mm = 160
    num_cols = len(rows[0])
    if num_cols > 0:
        col_width_mm = available_width_mm / num_cols
        for col in t.columns:
            col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    doc.add_paragraph()


def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Wastewater Treatment Plant Influent Flow and Load Assessment"
    today_str = date.today().strftime("%d %B %Y")

    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(
        doc,
        include_toc=True,
        include_lof=True,
        include_lot=True,
    )

    # 1 Introduction
    heading = doc.add_heading("1 Introduction", level=1)
    p = doc.add_paragraph(
        "This report presents an assessment of influent flow and load conditions for a municipal wastewater "
        "treatment plant (WWTP) based on four years of historical monitoring data and associated serviced "
        "population figures. The analysis derives representative average and peak design parameters and "
        "projects future flows and loads over a 20-year planning horizon using per capita rates and "
        "population growth forecasts."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "The outcomes are intended to support process engineering evaluations, capacity checks and "
        "conceptual upgrades of inlet works, biological treatment, clarification and sludge handling "
        "systems."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 2 Objectives
    doc.add_heading("2 Objectives", level=1)
    p = doc.add_paragraph(
        "The analysis has been structured to quantify historical influent behaviour, express it on a per "
        "capita basis and apply it to a population projection for planning purposes."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("The specific objectives are to:", style="Body Text")
    obj_list = [
        "Summarise historical influent flows and loads, including average annual and peak conditions.",
        "Derive per capita flow and load rates for biochemical oxygen demand (BOD), total suspended solids "
        "(TSS) and total phosphorus (TP).",
        "Characterise maximum month (30-day moving average) loads to reflect sustained peak operating "
        "conditions.",
        "Develop a 20-year serviced population projection and apply historical-average per capita rates to "
        "obtain projected flows and loads.",
        "Define design flow and load parameters at the end of the 20-year horizon and discuss implications "
        "for WWTP process capacity.",
    ]
    for item in obj_list:
        para = doc.add_paragraph(item, style="List Bullet")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 3 Description of Dataset
    doc.add_heading("3 Description of Dataset", level=1)
    p = doc.add_paragraph(
        "The assessment uses the workbook WWTP_medium.xlsx, which comprises a daily influent monitoring "
        "dataset and an annual serviced population series for the WWTP catchment."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "The “WWTP Influent Data” sheet provides a continuous daily time series from 1 January 2014 to "
        "31 December 2017 (1,461 records). Variables include influent flow rate (m³/d) and composite "
        "influent concentrations for BOD, TSS, total Kjeldahl nitrogen and phosphorus-related species. "
        "The “population” sheet reports the serviced population for each calendar year from 2014 to 2017."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4 Analysis
    doc.add_heading("4 Analysis", level=1)

    # 4.1 Historical flows
    doc.add_heading("4.1 Historical influent flows", level=2)
    p = doc.add_paragraph(
        "Average annual daily influent flows are very stable over 2014–2017, ranging from 28,658 to "
        "29,061 m³/d, with a historical mean of 28,815 m³/d. Maximum day flows are roughly 15–25% above "
        "the corresponding annual averages, with yearly maxima between 33,771 and 36,986 m³/d and an "
        "average of the annual maxima of 35,048 m³/d. This stability indicates a mature, largely domestic "
        "catchment with limited industrial or storm-driven variability."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("annual_flow_avg_max.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Annual average and maximum day influent flow (historical and overall average).",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "The ratio of maximum day to average day flow, derived as approximately 1.22 (35,048 / 28,815), "
        "is consistent across the four-year record. This peaking factor is an important parameter for "
        "checking hydraulic capacity of inlet works, screens, grit removal and downstream process units."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.2 Historical average daily loads
    doc.add_heading("4.2 Historical average daily loads", level=2)
    p = doc.add_paragraph(
        "Average annual daily loads for BOD, TSS and total phosphorus were calculated from measured "
        "concentrations and flows. Across 2014–2017, BOD loads range from approximately 5,413 to "
        "5,945 kg/d (historical mean 5,634 kg/d), TSS loads from 6,153 to 6,446 kg/d (mean 6,318 kg/d), "
        "and total phosphorus from 181 to 197 kg/d (mean 186 kg/d). Inter-annual variation is modest, "
        "generally less than 10%."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("annual_avg_loads_BOD_TSS_TP.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Average annual daily influent loads for BOD, TSS and total phosphorus.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "The relatively narrow spread of average loads suggests stable upstream contributions with no "
        "evidence of major step changes in industrial discharge. Slightly elevated BOD, TSS and "
        "phosphorus loads in 2016 are observable but do not indicate a structural change in catchment "
        "characteristics."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.3 Maximum month (30-day) loads
    doc.add_heading("4.3 Maximum month (30-day moving average) loads", level=2)
    p = doc.add_paragraph(
        "To capture sustained peak conditions relevant to biological process and sludge handling design, "
        "30-day moving average loads were calculated for each year and the maximum value identified."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("max_month_loads_BOD_TSS_TP.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Maximum month (30-day moving average) influent loads for BOD, TSS and total phosphorus.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "Maximum month BOD loads lie between 6,430 and 6,723 kg/d, with a historical mean of 6,593 kg/d, "
        "approximately 15–20% above the corresponding average-day loads. TSS maximum month loads are "
        "more strongly peaked, ranging from 7,565 to 9,093 kg/d (mean 8,159 kg/d), i.e. around 25–40% "
        "above average TSS loads. Maximum month phosphorus loads lie between 215 and 267 kg/d "
        "(mean 240 kg/d), about 25–30% above average-day TP."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "The highest maximum month TSS load occurs in 2015, indicating a period of sustained high-solids "
        "conditions that may have stressed primary clarification and sludge handling. Nevertheless, the "
        "data do not indicate chronic overload relative to current average conditions."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.4 Per capita flows
    doc.add_heading("4.4 Per capita flows", level=2)
    p = doc.add_paragraph(
        "Using the annual serviced population data, average-day and maximum-day flows were expressed on a "
        "per capita basis. Average per capita flows range from about 414 to 433 L/person/day, with a "
        "historical mean of 423.5 L/person/day. Maximum-day per capita flows range from about 490 to "
        "551 L/person/day with a historical mean of 515.3 L/person/day, consistent with the plant-wide "
        "peaking factor of approximately 1.22."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("annual_per_capita_flow_avg_max.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Annual average and maximum day influent flow per capita.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "The slight reduction in both average and maximum per capita flows in 2016 and 2017 could reflect "
        "incremental water efficiency improvements or small shifts in the serviced area. Overall, the "
        "per capita flow rates are typical for a combined domestic and minor commercial catchment."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.5 Per capita loads
    doc.add_heading("4.5 Per capita loads", level=2)
    p = doc.add_paragraph(
        "Per capita loads were derived by dividing calculated loads by the serviced population for each "
        "year. Average-day loads per capita are approximately 80–86 g BOD/person/day, 88.8–94.6 g "
        "TSS/person/day and 2.65–2.85 g total phosphorus/person/day, with historical means of "
        "82.8 g BOD/person/day, 92.9 g TSS/person/day and 2.7 g total phosphorus/person/day."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("annual_per_capita_avg_loads_BOD_TSS_TP.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Average annual daily influent loads per capita for BOD, TSS and total phosphorus.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "Maximum month (30-day) per capita loads are around 95.8–97.5 g BOD/person/day, "
        "109–136 g TSS/person/day and 3.2–3.9 g total phosphorus/person/day, with historical means of "
        "96.9 g BOD/person/day, 120 g TSS/person/day and 3.5 g total phosphorus/person/day."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("max_month_per_capita_loads_BOD_TSS_TP.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Maximum month (30-day moving average) influent loads per capita.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "These per capita values are consistent with typical domestic wastewater design ranges and "
        "indicate no extreme industrial contributions. The pronounced 2015 TSS maximum month peak at "
        "around 136 g/person/day suggests a temporary high-solids period that should be considered when "
        "evaluating clarifier and sludge system robustness."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.6 Historical-average per capita parameters (table)
    doc.add_heading("4.6 Historical-average per capita parameters", level=2)
    p = doc.add_paragraph(
        "Historical-average per capita flow and load parameters, averaged over the 2014–2017 period, are "
        "summarised in Table 1. These values form the basis for subsequent projections of flows and loads."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    per_capita_rows = [
        ["Parameter", "Value", "Units"],
        ["Average annual daily flow per capita", "423.5", "L/person/day"],
        ["Maximum day flow per capita", "515.3", "L/person/day"],
        ["Average day BOD load per capita", "82.8", "g/person/day"],
        ["Average day TSS load per capita", "92.9", "g/person/day"],
        ["Average day Total P load per capita", "2.7", "g/person/day"],
        ["Maximum month BOD load per capita (30-day avg)", "96.9", "g/person/day"],
        ["Maximum month TSS load per capita (30-day avg)", "119.9", "g/person/day"],
        ["Maximum month Total P load per capita (30-day avg)", "3.5", "g/person/day"],
    ]
    add_table(doc, "Historical-average per capita flow and load parameters.", per_capita_rows)

    # 4.7 Population growth and projection
    doc.add_heading("4.7 Population growth and projection", level=2)
    p = doc.add_paragraph(
        "The serviced population increased modestly from 66,880 in 2014 to 69,300 in 2017. A compound "
        "annual growth rate (CAGR) of approximately 1.19% per year was calculated from this period and "
        "used to project population to 2037. Under this assumption the serviced population is forecast "
        "to reach about 80,800 by 2030 and 87,800 by 2037, representing an increase of roughly 27% "
        "relative to 2017."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("population_projection_20yr.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Historical and projected serviced population over a 20-year horizon.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "The CAGR-based projection provides a smooth and defensible planning curve in the absence of "
        "specific information about step changes such as major developments or sewerage scheme "
        "extensions. Where such information exists, the projection should be refined accordingly."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.8 Projected flows based on per capita rates
    doc.add_heading("4.8 Projected flows using historical per capita rates", level=2)
    p = doc.add_paragraph(
        "Projected average-day and maximum-day flows were obtained by multiplying the projected "
        "population by the historical-average per capita flow rates from Table 1, namely "
        "423.5 L/person/day for average-day flow and 515.3 L/person/day for maximum-day flow. "
        "These rates maintain the observed peaking factor of approximately 1.22."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("historical_projected_flows_20yr.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Historical and projected average and maximum day influent flows.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "Under the adopted growth scenario, average-day flow increases from about 29,000 m³/d in 2017 to "
        "approximately 37,200 m³/d by 2037, while maximum-day flow rises from roughly 35,000 m³/d to "
        "around 45,300 m³/d. These flows should be compared with the hydraulic capacity of inlet "
        "structures, channels and critical treatment units for the design horizon."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.9 Projected average annual loads
    doc.add_heading("4.9 Projected average annual loads using per capita rates", level=2)
    p = doc.add_paragraph(
        "Average annual BOD, TSS and total phosphorus loads were projected using constant historical "
        "per capita average-day loads of 82.8 g BOD/person/day, 92.9 g TSS/person/day and "
        "2.7 g total phosphorus/person/day (Table 1), applied to the population forecast."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture("historical_projected_avg_loads_BOD_TSS_TP_20yr.png", width=Inches(5.5))
        add_caption(
            doc,
            "Figure",
            "Historical and projected average annual BOD, TSS and total phosphorus loads.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "Under these assumptions, average BOD load increases from around 5,800 kg/d in 2018 to "
        "approximately 7,300 kg/d by 2037. TSS loads increase from about 6,500 to 8,200 kg/d and "
        "total phosphorus from about 189 to 237 kg/d over the same period. This equates to a "
        "25–30% increase in organic and solids loading, with implications for aeration capacity, "
        "biological process sizing and sludge production."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.10 Projected maximum month loads
    doc.add_heading("4.10 Projected maximum month loads using per capita rates", level=2)
    p = doc.add_paragraph(
        "Maximum month (30-day average) projected loads were derived by applying constant historical "
        "maximum-month per capita loads of 96.9 g BOD/person/day, 120 g TSS/person/day and "
        "3.5 g total phosphorus/person/day (Table 1) to the population projection."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.add_picture(
            "historical_projected_max_month_loads_BOD_TSS_TP_20yr.png", width=Inches(5.5)
        )
        add_caption(
            doc,
            "Figure",
            "Historical and projected maximum month (30-day average) BOD, TSS and total phosphorus loads.",
        )
    except Exception:
        pass

    p = doc.add_paragraph(
        "By 2037 the projected maximum-month BOD load is approximately 8,500 kg/d, TSS about "
        "10,500 kg/d and total phosphorus around 310 kg/d. These sustained peak loads are roughly "
        "15–30% higher than projected average-day loads and are suitable for assessing biological "
        "process robustness, clarifier sizing and sludge handling infrastructure for extended wet "
        "weather or high-strength periods."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4.11 2037 design parameters
    doc.add_heading("4.11 Design parameters at 20-year horizon (2037)", level=2)
    p = doc.add_paragraph(
        "Design parameters at the end of the 20-year projection period (2037) were extracted from the "
        "projected series. These parameters represent the recommended basis for sizing new or upgraded "
        "treatment units, subject to local planning assumptions and risk appetite."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    design_2037_rows = [
        ["Parameter", "Value", "Units"],
        ["Average day flow", "37,200", "m³/d"],
        ["Maximum day flow", "45,260", "m³/d"],
        ["Average day BOD load", "7,270", "kg/d"],
        ["Average day TSS load", "8,160", "kg/d"],
        ["Average day Total P load", "240", "kg/d"],
        ["Maximum month BOD load (30-day avg)", "8,510", "kg/d"],
        ["Maximum month TSS load (30-day avg)", "10,530", "kg/d"],
        ["Maximum month Total P load (30-day avg)", "310", "kg/d"],
    ]
    add_table(
        doc,
        "Design flows and loads at the 2037 planning horizon based on projected population and "
        "historical-average per capita rates.",
        design_2037_rows,
    )

    # 5 Key Findings
    doc.add_heading("5 Key Findings", level=1)
    p = doc.add_paragraph(
        "Key findings from the historical analysis and 20-year projections can be summarised as follows."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    findings = [
        "Historical influent flows and loads are stable over 2014–2017, indicating a mature and "
        "largely domestic catchment with limited industrial variability.",
        "Average annual daily influent flow is approximately 28,800 m³/d, with an average annual "
        "maximum day of 35,000 m³/d, giving a stable peaking factor of about 1.22.",
        "Average-day per capita flows and loads (about 423.5 L/person/day, 82.8 g BOD/person/day, "
        "92.9 g TSS/person/day and 2.7 g total phosphorus/person/day) are consistent with standard "
        "municipal design values.",
        "Maximum month (30-day) loads increase per capita BOD, TSS and phosphorus by roughly 15–30% "
        "relative to average-day values, providing a suitable basis for design of sustained peak "
        "operating conditions.",
        "The serviced population is forecast to grow from 69,300 in 2017 to approximately 87,800 by "
        "2037 under a 1.19% per annum compound growth rate.",
        "Applying constant historical-average per capita rates yields 2037 design values of "
        "37,200 m³/d average-day flow, 45,260 m³/d maximum-day flow, average-day loads of about "
        "7,270 kg/d BOD, 8,160 kg/d TSS and 240 kg/d total phosphorus, and maximum-month loads of "
        "approximately 8,510 kg/d BOD, 10,530 kg/d TSS and 310 kg/d total phosphorus.",
    ]
    for item in findings:
        para = doc.add_paragraph(item, style="List Bullet")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 6 Recommendations
    doc.add_heading("6 Recommendations", level=1)
    p = doc.add_paragraph(
        "Based on the observed influent behaviour, per capita characteristics and projected growth, the "
        "following recommendations are made for process engineering and further investigation."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    recs_capacity = [
        "Use the 2037 design parameters in Table 2 as the primary basis for assessing hydraulic and "
        "treatment capacity of inlet works, biological reactors, clarifiers and sludge handling systems.",
        "Confirm that inlet structures and primary treatment units can reliably pass at least "
        "45,300 m³/d for maximum-day conditions without excessive headloss or bypass.",
        "Evaluate biological treatment and aeration capacity against projected maximum-month BOD loads "
        "of approximately 8,500 kg/d, allowing appropriate safety factors.",
        "Check clarifier and sludge handling performance under maximum-month TSS loads of around "
        "10,500 kg/d, particularly given the historically high 2015 solids peak.",
    ]
    for item in recs_capacity:
        para = doc.add_paragraph(item, style="List Bullet")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "Additional analyses and investigations would improve confidence in the projections and design "
        "parameters."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    further_analyses = [
        "Extend the influent monitoring period to include more recent years to confirm that per capita "
        "flows and loads remain stable, and to capture any recent water efficiency, industrial or "
        "catchment changes.",
        "Undertake seasonal and event-based analyses (e.g. wet- versus dry-weather periods) to better "
        "characterise infiltration/inflow and storm impacts on flow and solids loading.",
        "Review local planning information to refine the population projection, including any known "
        "major developments, sewerage extensions or industrial changes that could significantly alter "
        "flows or loads.",
        "Analyse nitrogen species in more detail (TKN, ammonia, nitrate and nitrite) to support "
        "biological nutrient removal design and to confirm consistency between organic and nutrient "
        "load trends.",
        "Consider sensitivity testing of per capita rates (for example, ±10–20%) and alternative growth "
        "rates to explore their effect on 20-year design parameters and to inform risk-based sizing.",
    ]
    for item in further_analyses:
        para = doc.add_paragraph(item, style="List Bullet")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()