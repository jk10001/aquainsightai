# filename: report.py
"""
Checklist (execution steps):
- Verify required chart image files and CSV summary files are present in the current directory.
- Load summary CSVs (without pandas) and build concise, report-ready tables (≤5 columns).
- Generate a DOCX with cover, updatable TOC/LoF/LoT fields, and consistent header/footer styling.
- Insert figures and tables with automatic numbering captions and reference them in the narrative.
- Save the final document as report.docx in the current directory.
"""

from __future__ import annotations

import csv
import os
from datetime import date
from typing import List, Dict, Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# ────────────────────────────────────────────────────────────────
#                       low-level helpers
# ────────────────────────────────────────────────────────────────
def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)  # A4
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> Paragraph:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")
    # Keep caption with the next block (table/picture)
    p.paragraph_format.keep_with_next = True
    return p


def add_header_footer(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        # Header
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        # Footer as a 2-col table (date left, page right)
        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        right_para = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            right_para.style = footer_style
        right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(right_para, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    doc.add_paragraph(title, style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(report_date, style="Subtitle").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document, include_toc: bool = True, include_lof: bool = True, include_lot: bool = True
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def _set_row_cant_split(row) -> None:
    """
    Prevent a Word table row from splitting across pages.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = trPr.find(qn("w:cantSplit"))
    if cant is None:
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    if not rows:
        return

    add_caption(doc, "Table", caption)

    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Repeat header row on each page (if table does break)
    t.rows[0].repeat_header = True

    # Prevent row splitting; improves readability when page breaks occur near a table
    for rr in t.rows:
        _set_row_cant_split(rr)

    # Fit to page width by distributing columns
    available_width_mm = 160
    num_cols = len(rows[0])
    if num_cols > 0:
        col_width_mm = available_width_mm / num_cols
        for col in t.columns:
            col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    doc.add_paragraph()


# ────────────────────────────────────────────────────────────────
#                     report-specific helpers
# ────────────────────────────────────────────────────────────────
def file_exists(path: str) -> bool:
    return os.path.isfile(path)


def fmt_int(x: Any) -> str:
    try:
        if x is None or x == "":
            return ""
        return f"{int(round(float(x))):,}"
    except Exception:
        return str(x)


def fmt_float(x: Any, ndp: int = 3) -> str:
    try:
        if x is None or x == "":
            return ""
        return f"{float(x):,.{ndp}f}"
    except Exception:
        return str(x)


def fmt_k(x: Any) -> str:
    try:
        if x is None or x == "":
            return ""
        return f"{float(x):.2f}"
    except Exception:
        return str(x)


def read_csv_dicts(path: str) -> List[Dict[str, str]]:
    if not file_exists(path):
        return []
    with open(path, "r", newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return [dict(r) for r in reader]


def add_body_paragraph(doc: Document, text: str, keep_with_next: bool = False) -> Paragraph:
    p = doc.add_paragraph(text, style="Body Text")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    return p


def add_spacer_paragraph(doc: Document, lines: int = 1) -> None:
    for _ in range(lines):
        doc.add_paragraph("", style="Body Text")


def add_heading(doc: Document, text: str, level: int = 1) -> Paragraph:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def safe_add_picture(doc: Document, img_path: str, width_mm: float = 160) -> bool:
    if not file_exists(img_path):
        p = doc.add_paragraph(
            f"Figure could not be inserted because the image file was not found: {img_path}",
            style="Body Text",
        )
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return False

    pic_par = doc.add_paragraph()
    run = pic_par.add_run()
    run.add_picture(img_path, width=Mm(width_mm))
    pic_par.paragraph_format.keep_with_next = True
    return True


# ────────────────────────────────────────────────────────────────
#                          MAIN report
# ────────────────────────────────────────────────────────────────
def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Catchment Rainfall and Reservoir Inflow Relationship (2015–2019)"
    today_str = date.today().strftime("%d %B %Y")

    # Cover & prelims
    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(doc, include_toc=True, include_lof=True, include_lot=True)

    # 1. Introduction
    add_heading(doc, "1. Introduction", level=1)
    add_body_paragraph(
        doc,
        "This report evaluates the relationship between catchment rainfall and inflow to major Melbourne Water supply reservoirs "
        "using daily time-series for 2015–2019. The primary focus is quantifying response timing and persistence using lagged rank "
        "correlation and an Antecedent Precipitation Index (API) model, and interpreting implications for short-term inflow predictability "
        "and operational planning.",
    )
    add_spacer_paragraph(doc)

    # 2. Objectives
    add_heading(doc, "2. Objectives", level=1)
    add_body_paragraph(
        doc,
        "The objectives of the analysis were to characterise whole-of-system storage behaviour and quantify rainfall–inflow coupling "
        "for reservoirs with applicable rainfall and inflow data, including response lags and effective catchment memory.",
    )
    doc.add_paragraph("The specific objectives were:", style="Body Text")
    doc.add_paragraph(
        "Summarise total stored water across all reservoirs relative to combined usable storage.", style="List Bullet"
    )
    doc.add_paragraph(
        "Visualise 2017 rainfall and inflow dynamics to qualitatively assess timing and seasonality.", style="List Bullet"
    )
    doc.add_paragraph(
        "Quantify lagged rainfall–inflow Spearman correlation for lags 0–14 days (rainfall leading).", style="List Bullet"
    )
    doc.add_paragraph(
        "Calibrate an API decay factor (k) to maximise same-day Spearman correlation between inflow and API.", style="List Bullet"
    )
    doc.add_paragraph(
        "Assess API-based predictability over 0–14 day lead times and summarise forecast skill.", style="List Bullet"
    )
    doc.add_paragraph(
        "Map reservoir locations to support interpretation of spatial rainfall variability and system operations.", style="List Bullet"
    )
    add_spacer_paragraph(doc)

    # 3. Description of Dataset
    add_heading(doc, "3. Description of Dataset", level=1)
    add_body_paragraph(
        doc,
        "The source dataset is an Excel workbook containing (i) a static asset register for major reservoirs "
        "(including usable volume and coordinates) and (ii) daily time-series (2015–2019) for catchment rainfall, "
        "streamflow into reservoirs (inflow), and storage volume/percent full. Rainfall and inflow time series were available for "
        "Upper Yarra, Thomson, O’Shannassy, and Maroondah. Storage volumes were available for a broader set of service and supply storages "
        "and were used to compute total system storage.",
    )
    add_spacer_paragraph(doc)

    # 4. Analysis
    add_heading(doc, "4. Analysis", level=1)

    # 4.1
    add_heading(doc, "4.1 System storage context (all reservoirs)", level=2)
    add_body_paragraph(
        doc,
        "Figure 1 summarises total stored water across all reservoirs as a stacked area chart, alongside the summed usable volume. "
        "This provides context for the rainfall–inflow relationships by indicating periods of drawdown and refill at a system level.",
        keep_with_next=True,
    )
    safe_add_picture(doc, "total_stored_vs_usable_volume.png", width_mm=160)
    add_caption(doc, "Figure", "Total stored water (all reservoirs) versus total usable volume.")
    add_spacer_paragraph(doc)

    add_body_paragraph(
        doc,
        "Across 2015–2019, total stored water generally remained below combined usable capacity, with a notable drawdown through 2018–2019 "
        "before partial recovery late-2019. Storage variability is dominated by Thomson (largest usable volume), with Upper Yarra and Cardinia "
        "as secondary contributors. The widening storage-to-capacity gap during 2018–2019 is consistent with reduced catchment yield and/or "
        "increased net offtakes during a dry period.",
    )
    add_spacer_paragraph(doc)

    # 4.2
    add_heading(doc, "4.2 Rainfall and inflow time-series behaviour (2017)", level=2)
    add_body_paragraph(
        doc,
        "To examine event timing and seasonal behaviour, Figure 2 presents 2017 daily catchment rainfall (bars) and reservoir inflow "
        "(line) using dual y-axes for each of the four catchments with complete rainfall and inflow records.",
        keep_with_next=True,
    )
    safe_add_picture(doc, "rainfall_inflow_2017_faceted.png", width_mm=160)
    add_caption(doc, "Figure", "2017 rainfall (catchment) and inflow (streamflow) time series by reservoir (dual y-axis).")
    add_spacer_paragraph(doc)

    add_body_paragraph(
        doc,
        "All catchments show a pronounced wet-season response in late winter to spring (approximately August to October), where "
        "frequent rainfall events coincide with sustained elevated inflows. Inflow peaks generally lag rainfall bursts by about 1–3 days, "
        "with elevated inflow persisting for weeks during the wet period. This sustained response indicates that baseflow and stored catchment "
        "moisture are important controls, rather than a purely event-driven runoff mechanism.",
    )
    add_spacer_paragraph(doc)

    # 4.3
    add_heading(doc, "4.3 Lagged rainfall–inflow correlation (0–14 days)", level=2)
    add_body_paragraph(
        doc,
        "Figure 3 quantifies the rank-based association between daily rainfall and inflow for lags of 0–14 days, where rainfall leads inflow. "
        "Spearman correlation is used to reduce sensitivity to non-linearity and extreme events.",
        keep_with_next=True,
    )
    safe_add_picture(doc, "lagged_rainfall_inflow_correlation_redo.png", width_mm=160)
    add_caption(doc, "Figure", "Lagged Spearman correlation between rainfall and inflow (rainfall leading, 0–14 days).")
    add_spacer_paragraph(doc)

    add_body_paragraph(
        doc,
        "Raw daily rainfall exhibits only weak-to-moderate correlation with inflow, with peak correlations typically occurring at lag 0–1 days "
        "and decaying rapidly thereafter. This behaviour is consistent with the physics of catchment response: inflow reflects both immediate runoff "
        "and delayed contributions (soil drainage, groundwater/baseflow), while rainfall is intermittent and spatially heterogeneous. "
        "The limited predictive skill of single-day rainfall motivates a state-based index such as API that explicitly represents antecedent wetness.",
    )
    add_spacer_paragraph(doc)

    # 4.4
    add_heading(doc, "4.4 API calibration (lag = 0 days, missing rainfall treated as 0)", level=2)
    add_body_paragraph(
        doc,
        "An Antecedent Precipitation Index (API) was computed for each catchment using a simple exponential decay formulation, treating missing "
        "rainfall as zero. The decay factor k was calibrated over 0.85–0.98 (approximately 0.01 steps) to maximise the absolute Spearman correlation "
        "between same-day inflow and API.",
        keep_with_next=True,
    )
    safe_add_picture(doc, "api_calibration_spearman.png", width_mm=160)
    add_caption(doc, "Figure", "API calibration: Spearman correlation versus API decay factor k (lag = 0 days).")
    add_spacer_paragraph(doc)

    add_body_paragraph(
        doc,
        "Figure 4 shows that correlation generally increases as k approaches 0.97–0.98, indicating that inflow is better explained by an index with "
        "substantial persistence. In process terms, the catchments behave as systems with significant storage and delayed release, where antecedent "
        "wetness sets the efficiency of rainfall-to-runoff transformation.",
    )
    add_spacer_paragraph(doc)

    calib_rows = read_csv_dicts("api_calibration_summary.csv")
    add_body_paragraph(
        doc,
        "Table 1 summarises the optimal k, corresponding maximum correlation, and inferred characteristic catchment memory time.",
        keep_with_next=True,
    )
    t1 = [["Reservoir", "Optimal k", "Spearman ρ (lag 0)", "Memory time (days)", "n (pairs)"]]
    for r in calib_rows:
        t1.append(
            [
                r.get("dam", ""),
                fmt_k(r.get("k_opt", "")),
                fmt_float(r.get("rho_at_k_opt", ""), 3),
                fmt_float(r.get("memory_time_days", ""), 1),
                fmt_int(r.get("n_pairs", "")),
            ]
        )
    add_table(doc, "API calibration summary (optimal k, same-day correlation, and memory time).", t1)

    add_body_paragraph(
        doc,
        "The calibrated k values are high (approximately 0.97–0.98), corresponding to characteristic memory times on the order of one to two months. "
        "O’Shannassy and Upper Yarra show the longest memory and strongest correlations, consistent with smoother hydrographs and a larger contribution "
        "from subsurface storage and baseflow. Maroondah exhibits the lowest peak correlation and shorter memory, indicating a relatively larger share "
        "of short-lived variability not captured by a single-parameter API model.",
    )
    add_spacer_paragraph(doc)

    # 4.5
    add_heading(doc, "4.5 API lead-time analysis (0–14 days) using optimal k", level=2)
    add_body_paragraph(
        doc,
        "Using the calibrated decay factors, Figure 5 evaluates how correlation changes when API is used to predict inflow at lead times "
        "of 0–14 days. This provides a practical indication of near-term forecast skill attributable to catchment wetness state.",
        keep_with_next=True,
    )
    safe_add_picture(doc, "api_lag_analysis.png", width_mm=160)
    add_caption(doc, "Figure", "API lead-time analysis: Spearman ρ versus lead time (0–14 days) at optimal k.")
    add_spacer_paragraph(doc)

    add_body_paragraph(
        doc,
        "Compared with Figure 3, API-based correlations are substantially higher and decay more gradually with lead time. This indicates that "
        "catchment wetness is a strong predictor of inflow persistence. In operational terms, a well-calibrated wetness index can support short-term "
        "yield outlooks and risk-based decision-making (e.g., supply planning and contingency triggers) even when event rainfall forecasts are uncertain.",
    )
    add_spacer_paragraph(doc)

    # Force Table 2 onto a new page to prevent it splitting (observed last-row spillover).
    doc.add_page_break()

    lag_rows = read_csv_dicts("api_lag_analysis_summary.csv")
    add_body_paragraph(
        doc,
        "Table 2 summarises key API lead-time performance metrics at the calibrated k values.",
        keep_with_next=True,
    )
    t2 = [["Reservoir", "Optimal k", "ρ (lag 0)", "ρ (lag 14)", "Horizon at ρ ≥ 0.75 (days)"]]
    for r in lag_rows:
        horizon = r.get("forecast_horizon_days_at_rho_ge_0_75", "")
        t2.append(
            [
                r.get("dam", ""),
                fmt_k(r.get("k_opt", "")),
                fmt_float(r.get("rho_lag0", ""), 3),
                fmt_float(r.get("rho_lag14", ""), 3),
                fmt_float(horizon, 0) if horizon not in (None, "", "nan", "NaN") else "",
            ]
        )
    add_table(doc, "API lead-time analysis summary at optimal k (selected performance metrics).", t2)

    add_body_paragraph(
        doc,
        "All catchments achieve maximum correlation at lag 0, indicating that API is most informative for same-day inflow. However, forecast skill "
        "differs markedly by catchment. O’Shannassy maintains strong correlation through 14 days, while Upper Yarra maintains high correlation out to "
        "approximately 10 days. Thomson decays more quickly, and Maroondah remains below the 0.75 threshold across all lead times, suggesting either "
        "greater flashiness, localised rainfall not captured by the catchment average rainfall series, and/or operational or measurement effects.",
    )
    add_spacer_paragraph(doc)

    # 4.6
    add_heading(doc, "4.6 Spatial context: reservoir locations", level=2)
    add_body_paragraph(
        doc,
        "Figure 6 maps reservoir locations based on the asset register coordinates. Spatial separation is relevant because storms can be localised; "
        "catchment-average rainfall signals may represent different synoptic regimes across the system, influencing comparative rainfall–inflow behaviour.",
        keep_with_next=True,
    )
    safe_add_picture(doc, "reservoir_locations_map.png", width_mm=160)
    add_caption(doc, "Figure", "Reservoir locations derived from the asset register (MGA Zone 55 converted to latitude/longitude).")
    add_spacer_paragraph(doc)

    # 5. Key Findings
    add_heading(doc, "5. Key Findings", level=1)
    add_body_paragraph(
        doc,
        "The key findings below synthesise results across Figures 1–6 and Tables 1–2, linking system storage behaviour with catchment response "
        "timing and persistence.",
    )
    doc.add_paragraph(
        "System storage remained below combined usable capacity throughout 2015–2019, with a material drawdown during 2018–2019 largely driven by Thomson (Figure 1).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "2017 time series show a wet-season regime with sustained inflow elevation beyond individual rainfall events, consistent with catchment wetting-up and baseflow contributions (Figure 2).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Raw rainfall is a weak predictor of inflow at daily resolution; peak Spearman correlation is modest and largely confined to lag 0–1 days (Figure 3).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "API calibration yields high optimal decay factors (k ≈ 0.97–0.98) and long catchment memory times (~33–49 days), indicating strong dependence on antecedent wetness (Figure 4, Table 1).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "API substantially improves correlation and retains predictive skill to 14 days, particularly for O’Shannassy and Upper Yarra, which exhibit the most persistent behaviour (Figure 5, Table 2).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Reservoirs are geographically dispersed, supporting the interpretation that spatial rainfall variability and differing catchment characteristics contribute to differences in correlation and persistence (Figure 6).",
        style="List Bullet",
    )
    add_spacer_paragraph(doc)

    # 6. Recommendations
    add_heading(doc, "6. Recommendations", level=1)
    add_body_paragraph(
        doc,
        "The analyses indicate that a persistent wetness state explains inflow substantially better than raw daily rainfall. The following recommendations "
        "support improved operational use and further technical refinement.",
    )
    doc.add_paragraph(
        "Extend API calibration beyond k = 0.98 (e.g., up to 0.995) for catchments that peaked at the current upper bound (Table 1), to confirm whether longer memory further improves performance.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Test alternative wetness indices (e.g., multi-store API, soil moisture accounting, or rainfall–runoff models with slow/fast stores) to better capture flashiness in catchments such as Maroondah (Figure 5, Table 2).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Incorporate seasonality explicitly (e.g., separate calibration for wet/dry seasons or temperature-based evapotranspiration proxies) to account for changing runoff efficiency.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Evaluate spatial rainfall representation (e.g., multiple gauges, gridded rainfall, or radar products) to reduce mismatch between catchment-average rainfall and true effective precipitation, particularly for localised storm events (Figure 6).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Assess operational influences and data artefacts (e.g., flow regulation, diversions, measurement timing) where lag-1 rainfall correlation peaks or where event responses appear inconsistent (Figure 3).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Develop a simple operational forecasting tool using calibrated API with lead-time performance benchmarks (Table 2), complemented by rainfall forecast scenarios for event-driven peaks.",
        style="List Bullet",
    )
    add_spacer_paragraph(doc)

    add_body_paragraph(
        doc,
        "Further analyses that would strengthen confidence include: quantile-based performance (high-flow versus low-flow regimes), event separation "
        "analysis to estimate unit hydrograph characteristics, and evaluation of non-stationarity (e.g., 2015–2016 versus 2018–2019) to link storage "
        "drawdown periods to shifts in catchment response and yield.",
    )

    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()