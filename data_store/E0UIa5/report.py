# filename: report.py
"""
Script to generate a comprehensive DOCX report analyzing anaerobic digester
instability at the Muscatine Wastewater Reclamation Facility (WRRF) for 2022.
"""

from datetime import date
from typing import List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Inches
from docx.text.paragraph import Paragraph


def configure_page(doc: Document) -> None:
    """Apply A4 size and 25 mm margins to every section."""
    for s in doc.sections:
        s.page_height = Mm(297)
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    """Add a Word field inside a run so it can be styled."""
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    """Automatic caption beginning with 'Figure/Table n – ...'."""
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(
    doc: Document, title: str, report_date: str | None = None
) -> None:
    """Add a header and a footer with date (left) and page number (right)."""
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        # Header
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        # Footer
        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Left cell - Date
        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        # Right cell - Page Number
        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style
        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    """Create a standalone title page followed by a page break."""
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    doc.add_paragraph(title, style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        report_date, style="Subtitle"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    """Insert TOC / List-of-Figures / List-of-Tables."""
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    """Create a formatted table with header repetition and auto-caption."""
    if not rows:
        return
    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Enable header row repetition across pages
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))

    # Prevent rows from splitting across pages
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

    available_width_mm = 160
    num_cols = len(rows[0])
    if num_cols > 0:
        col_width_mm = available_width_mm / num_cols
        for col in t.columns:
            col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()


def add_figure_image(doc: Document, image_path: str, caption_text: str) -> None:
    """Insert a centered figure image with standard width and auto-caption."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(5.8))
    add_caption(doc, "Figure", caption_text)
    doc.add_paragraph()


def generate_report() -> None:
    """Generate the complete Muscatine WRRF Digester Report."""
    doc = Document()
    configure_page(doc)

    report_title = "Evaluation of Anaerobic Digester Instability and Operational Performance at the Muscatine WRRF"
    report_date = "24 May 2024"

    # Cover & Prelims
    add_title_page(doc, report_title, report_date)
    add_header_footer(doc, report_title, report_date)
    add_front_matter(
        doc,
        include_toc=True,
        include_lof=True,
        include_lot=True,
    )

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    
    p = doc.add_paragraph(
        "The Muscatine Wastewater Reclamation Facility (WRRF) relies on a dual mesophilic anaerobic digestion system "
        "to stabilize primary sludge (PS), thickened waste activated sludge (TWAS), and high-strength waste (HSW) streams. "
        "The facility operates two primary anaerobic digesters, designated as Digester 1 and Digester 2, each having an active "
        "liquid volume of 1,625 m³, yielding a combined digestion volume of 3,250 m³. Under baseline operating procedures, "
        "the municipal sludge loads (TWAS and PS) are distributed equally between the two digesters. However, co-digestion "
        "substrates—specifically liquid HSW and Fats, Oils, and Grease (FOG)—are co-fed into the digesters based on operational "
        "routing decisions."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "Co-digestion of high-strength organic wastes offers significant benefits, including enhanced biogas generation, "
        "improved renewable energy recovery, and additional tipping fee revenue. Nevertheless, co-digestion introduces "
        "biochemical vulnerabilities. Highly biodegradable organic streams, such as food waste, grease trap waste, and ethanol "
        "byproducts, rapidly hydrolyze into volatile fatty acids (VFAs). If organic loading rates exceed the kinetic capacity "
        "of syntrophic acetogens and methanogenic archaea, VFAs accumulate, consume natural bicarbonate buffering, depress pH, "
        "and ultimately result in severe process souring."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "During calendar year 2022, operational monitoring revealed pronounced performance disparities between Digester 1 and "
        "Digester 2, culminating in severe souring events in Digester 1 during late September and October. This report presents "
        "a rigorous process engineering analysis of the 2022 daily monitoring dataset to diagnose the root causes of digester "
        "instability, evaluate operational feed routing practices, quantify thermodynamic and biochemical impacts, and establish "
        "robust corrective measures."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 2. Objectives
    doc.add_heading("2. Objectives", level=1)

    p = doc.add_paragraph(
        "The primary objective of this investigation is to systematically evaluate the operational stability and biochemical "
        "performance of Digester 1 and Digester 2 at the Muscatine WRRF throughout calendar year 2022. Specific engineering "
        "objectives include:"
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p1 = doc.add_paragraph(style="List Bullet")
    p1.add_run("Chronologically map process instability events using established indicator thresholds (VFA/Alkalinity ratio, pH, total alkalinity, OLR, and HRT).")
    
    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run("Evaluate the impact of asymmetric High-Strength Waste (HSW) routing on volumetric loading rates, hydraulic retention times, and buffering capacity.")

    p3 = doc.add_paragraph(style="List Bullet")
    p3.add_run("Analyze the interconnections between feed distribution changes, acid accumulation, alkalinity depletion, and plantwide biogas yield.")

    p4 = doc.add_paragraph(style="List Bullet")
    p4.add_run("Formulate actionable engineering recommendations, process control caps, and operational safeguards to prevent recurrent souring episodes.")

    # 3. Description of Dataset
    doc.add_heading("3. Description of Dataset", level=1)

    p = doc.add_paragraph(
        "The evaluation is based on daily operational and monitoring data collected at the Muscatine WRRF from 1 January 2022 "
        "through 31 December 2022 (365 daily observations). The dataset encompasses feed stream volumetric flows (TWAS, PS, HSW, FOG), "
        "feed characteristics (TWAS-VS%, PS-VS%, HSW-VS%, HSW-COD), HSW feed split percentages between digesters, individual digester "
        "operating parameters (temperature, pH, alkalinity, VFA, calculated VFA/Alkalinity ratio), and facility biogas production. "
        "Plant influent flow during 2022 averaged 11,337 m³/d, while daily biogas generation averaged 4,716 m³/d. Total HSW receiving "
        "averaged 56.5 m³/d, with an average volatile solids content of 7.26% and COD concentration of 158,210 mg/L."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 4. Analysis
    doc.add_heading("4. Analysis", level=1)

    p = doc.add_paragraph(
        "The process engineering analysis evaluates digester performance across key biochemical, physical, and operational metrics. "
        "By synthesizing time-series monitoring data, the physical mechanisms driving system instability were identified and quantified."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_heading("4.1 Acidification and Volatile Fatty Acid Dynamics", level=2)

    p = doc.add_paragraph(
        "The ratio of Volatile Fatty Acids (VFA) to total alkalinity (expressed as mg/L VFA as acetic acid divided by mg/L alkalinity "
        "as CaCO3) serves as the primary early-warning indicator for anaerobic digestion stability. A VFA/Alkalinity ratio below 0.30 "
        "indicates robust methanogenic equilibrium. Ratios between 0.30 and 0.40 signal process stress requiring operational intervention, "
        "while ratios exceeding 0.40 denote severe organic overload, acute acid accumulation, and imminent souring."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "As illustrated in Figure 1, Digester 1 experienced extreme instability during late September and October 2022. The VFA/Alkalinity "
        "ratio in Digester 1 surged above the 0.40 upset threshold continuously from 21 September through 20 October 2022, reaching an "
        "unprecedented peak of 0.930 on 30 September 2022 (with secondary peaks of 0.735 on 29 September, 0.675 on 13 October, and 0.631 on "
        "19 October). This sustained elevation confirms an unmitigated souring event where volatile acid production completely overwhelmed "
        "methanogenic utilization. Conversely, Digester 2 maintained high stability throughout this period, operating with VFA/Alkalinity "
        "ratios between 0.12 and 0.22, though it experienced brief transient spikes above 0.40 in late January (peak 0.520 on 26 January) "
        "and early October (0.455 on 7 October)."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_vfa_alk_ratio.png",
        "Digester VFA to alkalinity ratio time series for Digester 1 and Digester 2 during 2022 showing critical threshold exceedances."
    )

    doc.add_heading("4.2 pH and Bicarbonate Alkalinity Buffering", level=2)

    p = doc.add_paragraph(
        "The pH of an anaerobic digester is buffered primarily by the carbonic acid-bicarbonate system. Methanogenic archaea operate "
        "within an optimal pH window of 6.8 to 7.2. Below a pH of 6.5, methanogenic activity drops dramatically, leading to an irreversible "
        "acidification loop if feed loading is maintained."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "Figure 2 illustrates the pH response of both digesters. In Digester 1, pH dropped below the typical 6.8 lower boundary repeatedly "
        "during the late September to October upset window, reaching critical minimums of 6.50 on 30 September, 6.50 on 13 October, and 6.43 "
        "on 20 October 2022. This severe pH depression correlates directly with the VFA/Alkalinity spikes shown in Figure 1."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_ph_time_series.png",
        "Digester pH time series for Digester 1 and Digester 2 showing severe pH depression in Digester 1 during September–October 2022."
    )

    p = doc.add_paragraph(
        "The underlying cause of pH failure is the exhaustion of bicarbonate buffering capacity. Anaerobic digesters require total "
        "alkalinity above 3,000 mg/L as CaCO3 (ideally 4,000 to 5,000 mg/L) to neutralize metabolic acids. Figure 3 demonstrates that "
        "Digester 1 alkalinity collapsed from ~5,500 mg/L in early September to below 3,000 mg/L on 24 September, hitting a minimum of "
        "1,965 mg/L on 20 October 2022. Digester 2 maintained robust alkalinity (4,000 to 5,700 mg/L) during this window. However, both "
        "digesters exhibited a shared alkalinity depression in late May (dropping to ~2,222 mg/L in Digester 2 and ~2,935 mg/L in Digester 1), "
        "and both experienced a steady downward trend in December (2,700 to 3,400 mg/L), indicating systemic winter buffer erosion."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_alkalinity_time_series.png",
        "Digester alkalinity time series for Digester 1 and Digester 2 detailing total buffering reserve depletion."
    )

    doc.add_heading("4.3 Volumetric Loading and Retention Time Asymmetry", level=2)

    p = doc.add_paragraph(
        "Digester hydraulic retention time (HRT) and volatile solids organic loading rate (VS OLR) are direct functions of active volume "
        "and feed flow rate. For mesophilic digesters treating mixed sludges and HSW, an HRT of 15 to 25 days and a maximum OLR of "
        "3.2 kg VS/m³-d represent conservative design boundaries."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "Figure 4 presents the calculated individual digester HRT based on active volumes of 1,625 m³ each. From July through September 2022, "
        "Digester 1 HRT was suppressed below the 15-day reference limit, frequently operating between 9 and 13 days. Simultaneously, "
        "Digester 2 HRT elevated to 30–55 days. This severe hydraulic disparity was caused directly by operational routing of High-Strength Waste."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_hrt_time_series.png",
        "Calculated Hydraulic Retention Time (HRT) time series by digester illustrating sustained hydraulic overloading of Digester 1."
    )

    p = doc.add_paragraph(
        "The HSW split fraction depicted in Figure 6 reveals that from 8 July through 29 September 2022, approximately 100% of all incoming HSW "
        "was preferentially routed into Digester 1, while municipal sludges (TWAS and PS) were split 50/50. As a result, Digester 1 received "
        "an extreme concentration of highly biodegradable organic load."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_hsw_split_fraction.png",
        "High-Strength Waste (HSW) split fraction between Digester 1 and Digester 2 showing prolonged 100% allocation to Digester 1."
    )

    p = doc.add_paragraph(
        "Figure 5 illustrates the 7-day rolling average Volatile Solids OLR. Driven by the 100% HSW feed allocation, Digester 1 OLR surged "
        "above the 3.2 kg VS/m³-d upper threshold from mid-July through early October, reaching a peak of 4.57 kg VS/m³-d on 29 September 2022. "
        "During this same period, Digester 2 OLR was starved down to 0.5–1.0 kg VS/m³-d. In early October, plant personnel abruptly flipped "
        "the HSW feed 100% to Digester 2. This resulted in an immediate OLR crossover, subjecting Digester 2 to a step OLR shock of "
        "3.0 kg VS/m³-d while Digester 1 OLR collapsed to 0.46 kg VS/m³-d. This uncoordinated redistribution imposed shock-loading "
        "stresses on Digester 2 and disrupted biomass acclimation."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_vs_loading_7d_avg.png",
        "Seven-day rolling average Volatile Solids Organic Loading Rate (VS OLR) highlighting prolonged organic overloading of Digester 1."
    )

    doc.add_heading("4.4 Biogas Production and Temperature Stability", level=2)

    p = doc.add_paragraph(
        "Facility biogas production averaged 4,716 m³/d during 2022 but exhibited extreme daily volatility ranging from 155 m³/d to "
        "9,350 m³/d, as shown in Figure 7. During the initial phase of Digester 1 overloading in late September, biogas production spiked "
        "to 5,600–7,700 m³/d. This temporary gas surge reflected rapid acidogenesis and acetogenesis of readily soluble HSW organics. "
        "However, because methanogenesis could not keep pace with acid generation, VFAs accumulated rapidly, consuming alkalinity and "
        "initiating the souring cascade. The single-day drop to 155 m³/d on 6 November 2022 represents an instrumentation or flaring outage "
        "rather than a biological cessation."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_biogas_time_series.png",
        "Daily biogas production time series showing high process volatility and transient peak yields during organic overloading."
    )

    p = doc.add_paragraph(
        "As shown in Figure 8, operating temperatures in both digesters were generally well maintained within the target mesophilic range "
        "of 34 °C to 36 °C (annual average ~35.5 °C). Digester 1 experienced transient low-temperature anomalies in March (33.3 °C) and "
        "sharp single-day drops in August (down to 31.1 °C on 4 August and 31.3 °C on 31 August). While temperature instability was not the "
        "primary driver of the September souring, cold shocks impair methanogenic kinetics and reduce acid-clearing capacity."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_figure_image(
        doc,
        "muscatine_wrrf_digesters_2022_temperature_time_series.png",
        "Digester operating temperature time series confirming baseline mesophilic stability punctuated by isolated cold-shock anomalies."
    )

    doc.add_heading("4.5 Chronology of Major Instability Events", level=2)

    p = doc.add_paragraph(
        "A threshold-based event detection analysis was conducted across the 2022 dataset. Instability events were classified based on "
        "VFA/Alk > 0.40, pH < 6.8, alkalinity < 3,000 mg/L as CaCO3, and 7-day VS OLR > 3.2 kg VS/m³-d. Table 1 summarizes the primary "
        "instability episodes identified during calendar year 2022."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    table_data = [
        ["Digester", "Event Period", "Duration", "Key Metrics", "Process Impact & Diagnostic Notes"],
        [
            "Digester 1",
            "2022-07-20 to 2022-07-24",
            "5 days",
            "Max OLR: 3.69 kgVS/m³-d\nHSW Split: 100% Dig 1",
            "Early loading imbalance caused by total HSW allocation to Digester 1; pre-conditioned system for subsequent instability."
        ],
        [
            "Digester 1",
            "2022-09-20 to 2022-10-03",
            "14 days",
            "Max VFA/Alk: 0.930\nPeak OLR: 4.57 kgVS/m³-d",
            "Severe souring episode: 8 days VFA/Alk > 0.40, min pH 6.50, alkalinity collapse to 2,555 mg/L as CaCO3. Primary operational failure."
        ],
        [
            "Digester 1",
            "2022-10-13 to 2022-10-21",
            "6 days*",
            "Min pH: 6.43\nMin Alk: 1,965 mg/L",
            "Secondary souring window following feed redirection: critical buffering exhaustion, requiring prolonged recovery and load shed."
        ],
        [
            "Digester 2",
            "2022-01-24 to 2022-01-27",
            "4 days",
            "Max VFA/Alk: 0.520\nAlk: > 5,700 mg/L",
            "Transient VFA accumulation during high winter organic loading; system remained buffered by elevated background alkalinity."
        ],
        [
            "Digester 2",
            "2022-10-07",
            "1 day",
            "Max VFA/Alk: 0.455\nHSW Split: 100% Dig 2",
            "Abrupt step-loading shock response following emergency 100% HSW feed redirection from Digester 1 to Digester 2."
        ],
        [
            "Digester 2",
            "2022-12-02 to 2022-12-16",
            "15 days",
            "Min Alk: 2,587 mg/L\nMin pH: 6.69",
            "Systemic winter buffering depletion (13 days Alk < 3,000 mg/L) with mild pH depression under steady baseline municipal loading."
        ],
    ]

    add_table(doc, "Summary chronology of primary digester instability events during calendar year 2022.", table_data)

    p = doc.add_paragraph("Note: *Denotes cumulative days across two closely spaced sub-events (13–15 October and 19–21 October 2022).")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 5. Key Findings
    doc.add_heading("5. Key Findings", level=1)

    p = doc.add_paragraph(
        "Synthesizing the analytical results across all operational streams reveals several fundamental process insights regarding "
        "the Muscatine WRRF anaerobic digestion system:"
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p1 = doc.add_paragraph(style="List Bullet")
    p1.add_run("Uncontrolled Asymmetric Feed Routing as Root Cause: The primary driver of severe souring in Digester 1 was the prolonged, 100% preferential routing of High-Strength Waste (HSW) from July through September 2022. This asymmetric loading violated basic mass-balance principles.")

    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run("Organic Overloading and Retention Time Suppression: Routing total HSW into a single 1,625 m³ digester elevated 7-day average OLRs to 4.57 kg VS/m³-d (43% above the conservative 3.2 threshold) and suppressed HRT to 9–13 days (below the 15-day minimum requirement), causing methanogenic washout.")

    p3 = doc.add_paragraph(style="List Bullet")
    p3.add_run("Sequential Acidification and Buffer Collapse: Acidification followed a classic biochemical cascade: rapid HSW hydrolysis caused VFAs to accumulate (>2,500 mg/L), which depleted bicarbonate alkalinity (dropping from 5,500 to 1,965 mg/L), driving VFA/Alk ratios to 0.930 and depressing pH to 6.43.")

    p4 = doc.add_paragraph(style="List Bullet")
    p4.add_run("Destabilizing Shock Loading from Emergency Feed Flips: In early October, operational attempts to correct Digester 1 souring by abruptly re-routing 100% of HSW to Digester 2 caused an immediate step-loading shock (OLR surging from 0.5 to 3.0 kg VS/m³-d), destabilizing Digester 2.")

    p5 = doc.add_paragraph(style="List Bullet")
    p5.add_run("Misleading Early Biogas Signals: Biogas production peaked during the initial phase of overloading due to rapid acidogenesis of HSW. Relying solely on biogas volume as a health metric masked severe underlying alkalinity depletion.")

    p6 = doc.add_paragraph(style="List Bullet")
    p6.add_run("Systemic Winter Buffering Vulnerability: Both digesters exhibited declining alkalinity during December (<3,000 mg/L), indicating elevated vulnerability to acid upsets during cold-weather operations.")

    # 6. Recommendations
    doc.add_heading("6. Recommendations", level=1)

    p = doc.add_paragraph(
        "To prevent future souring events, ensure dual-digester stability, and optimize co-digestion energy recovery, the following "
        "engineering and operational measures are recommended:"
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_heading("6.1 Immediate Operational Controls and Limits", level=2)

    p = doc.add_paragraph(
        "1. Automated 50/50 HSW Feed Splitting: Implement automated flow-proportional control on the HSW receiving feed manifold to maintain an equal 50/50 volumetric and organic load distribution between Digester 1 and Digester 2 at all times."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "2. Maximum Volatile Solids OLR Cap: Enforce a strict upper cap of 3.0 kg VS/m³-d for individual digester 7-day rolling average OLR. If total incoming HSW exceeds this threshold, excess HSW must be diverted to temporary storage or offsite disposal."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "3. Minimum HRT Safeguard: Restrict total daily liquid feed volumes (TWAS + PS + HSW) to ensure individual digester HRT never drops below 15 days (maximum liquid feed rate of 108 m³/d per digester)."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "4. Alkalinity Buffer Management and Dosing Plan: Establish a minimum alkalinity threshold of 3,500 mg/L as CaCO3. Procure a chemical dosing system (e.g., sodium bicarbonate or liquid lime) to automatically dose buffer whenever digester alkalinity drops below 3,000 mg/L or VFA/Alk exceeds 0.30."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_heading("6.2 Further Data Analysis and Modeling", level=2)

    p = doc.add_paragraph(
        "1. Specific Biogas Yield Analysis: Calculate daily specific biogas yield (m³ biogas / kg VS destroyed) to accurately differentiate between true methanogenic conversion efficiency and transient acidogenesis gas spikes."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "2. Volatile Solids Destruction Mass Balance: Perform rigorous mass balances on total and volatile solids across the primary clarification, thickening, and digestion processes to evaluate organic conversion rates."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "3. Process Kinetic Modeling: Develop an Anaerobic Digestion Model No. 1 (ADM1) or simplified kinetic model to simulate dynamic response times under step HSW load changes."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_heading("6.3 Field Investigations and Process Testing", level=2)

    p = doc.add_paragraph(
        "1. HSW Characterization and Toxicity Screening: Implement routine laboratory testing of incoming HSW deliveries for soluble COD, total VFAs, pH, free ammonia, and long-chain fatty acids (LCFA) to screen for toxic loads prior to digester injection."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "2. Continuous Online Monitoring Installation: Install online pH, oxidation-reduction potential (ORP), and continuous biogas methane/CO2 analyzer probes to provide real-time early warning of acidification."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = doc.add_paragraph(
        "3. Tracer and Mixing Profiling Studies: Perform lithium chloride or fluoride tracer tests on Digester 1 and Digester 2 to determine actual hydraulic active volume, identify short-circuiting, and evaluate mixing efficiency."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save Document
    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()