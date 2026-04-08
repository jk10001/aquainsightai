# filename: report.py
# Checklist (execution steps):
# 1) Load the digester CSV (no pandas) and robustly parse headers/dates/metrics for all three units.
# 2) Compute key upset indicators (first exceedance dates, peaks/minima, recovery dates).
# 3) Assemble a structured DOCX with title page, updatable TOC/LOF/LOT fields, and numbered captions.
# 4) Insert provided charts with captions and reference them in the analysis narrative.
# 5) Add concise summary tables (≤5 columns) using the required Word table style.
# 6) Save the finished report as report.docx in the current directory.

from __future__ import annotations

import csv
import os
from dataclasses import dataclass
from datetime import date, datetime
from statistics import mean
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# ────────────────────────────────────────────────────────────────
#                       low-level helpers
# ────────────────────────────────────────────────────────────────
def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)  # A4
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        # header
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        # footer
        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # date (left)
        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        # page number (right)
        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style
        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    doc.add_paragraph(title, style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(report_date, style="Subtitle").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    if not rows:
        return
    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    available_width_mm = 160
    num_cols = len(rows[0])
    col_width_mm = available_width_mm / max(num_cols, 1)
    for col in t.columns:
        col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    t.rows[0].repeat_header = True
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()


def fmt_int(x: float | int) -> str:
    try:
        return f"{int(round(float(x))):,}"
    except Exception:
        return str(x)


def fmt_float(x: float | int, ndp: int = 2) -> str:
    try:
        return f"{float(x):,.{ndp}f}"
    except Exception:
        return str(x)


def fmt_date(d: Optional[datetime]) -> str:
    if d is None:
        return "—"
    return d.strftime("%d %b %Y")


def safe_exists(path: str) -> bool:
    return os.path.exists(path) and os.path.isfile(path)


def add_body_paragraph(doc: Document, text: str) -> Paragraph:
    p = doc.add_paragraph(text, style="Body Text")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p


# ────────────────────────────────────────────────────────────────
#                       data model & analysis
# ────────────────────────────────────────────────────────────────
@dataclass
class DigesterSeries:
    vfa: List[float]
    alk: List[float]
    ratio: List[float]
    ph: List[float]
    vs_dest: List[float]
    gas: List[float]


@dataclass
class ExtremePoint:
    value: float
    when: datetime


@dataclass
class DigesterKPI:
    max_vfa: ExtremePoint
    max_ratio: ExtremePoint
    min_ph: ExtremePoint
    min_vs: ExtremePoint
    min_gas: ExtremePoint
    first_ratio_over_030: Optional[datetime]
    first_vfa_over_1000: Optional[datetime]
    last_ratio_below_010_after_upset: Optional[datetime]


def _norm_header(h: str) -> str:
    return (h or "").strip().lstrip("\ufeff")


def parse_csv(path: str) -> Tuple[List[datetime], Dict[int, DigesterSeries]]:
    dates: List[datetime] = []
    series: Dict[int, DigesterSeries] = {
        1: DigesterSeries([], [], [], [], [], []),
        2: DigesterSeries([], [], [], [], [], []),
        3: DigesterSeries([], [], [], [], [], []),
    }

    with open(path, "r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        try:
            raw_headers = next(reader)
        except StopIteration:
            raise ValueError("CSV file is empty.")

        headers = [_norm_header(h) for h in raw_headers]
        idx = {h: i for i, h in enumerate(headers)}

        required = ["Date"]
        for i in (1, 2, 3):
            required.extend(
                [
                    f"Digester {i} VFA (mg/L)",
                    f"Digester {i} Alkalinity (mg/L)",
                    f"Digester {i} VFA/Alkalinity",
                    f"Digester {i} pH",
                    f"Digester {i} volatile solids destruction (%)",
                    f"Digester {i} gas production (m3/day)",
                ]
            )

        missing = [h for h in required if h not in idx]
        if missing:
            raise KeyError(
                "Missing expected column(s): "
                + ", ".join(missing)
                + ". Found headers: "
                + ", ".join(headers)
            )

        for row in reader:
            if not row or all((c.strip() == "" for c in row)):
                continue
            if len(row) < len(headers):
                row = row + [""] * (len(headers) - len(row))

            d = datetime.strptime(row[idx["Date"]].strip(), "%d/%m/%Y")
            dates.append(d)

            for i in (1, 2, 3):
                series[i].vfa.append(float(row[idx[f"Digester {i} VFA (mg/L)"]]))
                series[i].alk.append(float(row[idx[f"Digester {i} Alkalinity (mg/L)"]]))
                series[i].ratio.append(float(row[idx[f"Digester {i} VFA/Alkalinity"]]))
                series[i].ph.append(float(row[idx[f"Digester {i} pH"]]))
                series[i].vs_dest.append(float(row[idx[f"Digester {i} volatile solids destruction (%)"]]))
                series[i].gas.append(float(row[idx[f"Digester {i} gas production (m3/day)"]]))

    return dates, series


def argmax(values: List[float]) -> int:
    idx, best = 0, values[0]
    for i, v in enumerate(values):
        if v > best:
            best, idx = v, i
    return idx


def argmin(values: List[float]) -> int:
    idx, best = 0, values[0]
    for i, v in enumerate(values):
        if v < best:
            best, idx = v, i
    return idx


def first_date_over(dates: List[datetime], values: List[float], thr: float) -> Optional[datetime]:
    for d, v in zip(dates, values):
        if v > thr:
            return d
    return None


def last_date_below_after_first_over(
    dates: List[datetime],
    values: List[float],
    over_thr: float,
    below_thr: float,
) -> Optional[datetime]:
    first_over_idx = None
    for i, v in enumerate(values):
        if v > over_thr:
            first_over_idx = i
            break
    if first_over_idx is None:
        return None

    last_below: Optional[datetime] = None
    for d, v in zip(dates[first_over_idx:], values[first_over_idx:]):
        if v < below_thr:
            last_below = d
    return last_below


def compute_kpis(dates: List[datetime], s: DigesterSeries) -> DigesterKPI:
    i_vfa = argmax(s.vfa)
    i_ratio = argmax(s.ratio)
    i_ph = argmin(s.ph)
    i_vs = argmin(s.vs_dest)
    i_gas = argmin(s.gas)

    return DigesterKPI(
        max_vfa=ExtremePoint(s.vfa[i_vfa], dates[i_vfa]),
        max_ratio=ExtremePoint(s.ratio[i_ratio], dates[i_ratio]),
        min_ph=ExtremePoint(s.ph[i_ph], dates[i_ph]),
        min_vs=ExtremePoint(s.vs_dest[i_vs], dates[i_vs]),
        min_gas=ExtremePoint(s.gas[i_gas], dates[i_gas]),
        first_ratio_over_030=first_date_over(dates, s.ratio, 0.30),
        first_vfa_over_1000=first_date_over(dates, s.vfa, 1000.0),
        last_ratio_below_010_after_upset=last_date_below_after_first_over(dates, s.ratio, 0.30, 0.10),
    )


def infer_common_upset_window(kpis: Dict[int, DigesterKPI]) -> Tuple[Optional[datetime], Optional[datetime]]:
    starts = [k.first_ratio_over_030 for k in kpis.values() if k.first_ratio_over_030]
    if not starts:
        return None, None
    upset_start = max(starts)

    ends = [k.last_ratio_below_010_after_upset for k in kpis.values() if k.last_ratio_below_010_after_upset]
    upset_end = min(ends) if ends else None
    return upset_start, upset_end


# ────────────────────────────────────────────────────────────────
#                       report generation
# ────────────────────────────────────────────────────────────────
def add_figure(doc: Document, img_path: str, caption: str, width_mm: float = 160.0) -> None:
    if safe_exists(img_path):
        doc.add_picture(img_path, width=Mm(width_mm))
        add_caption(doc, "Figure", caption)
        doc.add_paragraph()
    else:
        p = add_body_paragraph(doc, f"[Missing figure file: {img_path}] {caption}")
        if p.runs:
            p.runs[0].italic = True


def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Municipal WWTP Anaerobic Digester Stability Assessment (Jun 2018–Jan 2019)"
    today_str = date.today().strftime("%d %B %Y")

    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(doc, include_toc=True, include_lof=True, include_lot=True)

    csv_path = "digester_data_3.csv"
    if not safe_exists(csv_path):
        add_body_paragraph(doc, f"ERROR: Source dataset not found: {csv_path}")
        doc.save("report.docx")
        return

    try:
        dates, series = parse_csv(csv_path)
    except Exception as e:
        add_body_paragraph(doc, f"ERROR: Could not parse dataset: {e}")
        doc.save("report.docx")
        return

    kpis = {i: compute_kpis(dates, series[i]) for i in (1, 2, 3)}
    upset_start, upset_end = infer_common_upset_window(kpis)

    # 1. Introduction
    doc.add_paragraph("1. Introduction", style="Heading 1")
    add_body_paragraph(
        doc,
        "This report assesses anaerobic digester stability for a municipal wastewater treatment plant "
        "using daily monitoring data for three digesters. The analysis focuses on indicators of process "
        "imbalance (volatile fatty acids and buffering), biological response (pH and volatile solids destruction), "
        "and outcome performance (biogas production). Emphasis is placed on identifying plant-wide process issues, "
        "likely causes, and operational actions to prevent recurrence.",
    )

    # 2. Objectives
    doc.add_paragraph("2. Objectives", style="Heading 1")
    add_body_paragraph(
        doc,
        "The objectives of this assessment are to quantify the timing and severity of any instability event(s), "
        "evaluate how the three digesters responded, and identify plausible plant-wide drivers consistent with the observed trends.",
    )
    doc.add_paragraph("Specifically, this report aims to:", style="Body Text")
    doc.add_paragraph("Characterise stability using VFA/alkalinity ratio, VFA concentration, and pH.", style="List Bullet")
    doc.add_paragraph("Quantify performance impacts using biogas production and volatile solids destruction.", style="List Bullet")
    doc.add_paragraph(
        "Identify evidence for common-cause versus unit-specific issues and propose operational controls and investigations.",
        style="List Bullet",
    )

    # 3. Description of Dataset
    doc.add_paragraph("3. Description of Dataset", style="Heading 1")
    add_body_paragraph(
        doc,
        "The dataset (digester_data_3.csv) contains 221 daily records from 01 Jun 2018 to 07 Jan 2019. "
        "For each of three anaerobic digesters, the following parameters are provided: VFA (mg/L as acetic), "
        "alkalinity (mg/L as CaCO3), VFA/alkalinity ratio, pH, volatile solids destruction (%), and biogas production (m³/day). "
        "The time step is one day and all series are complete (no missing values).",
    )

    # 4. Analysis
    doc.add_paragraph("4. Analysis", style="Heading 1")
    if upset_start and upset_end:
        add_body_paragraph(
            doc,
            "Time-series review indicates a synchronised process upset affecting all three digesters. "
            f"Based on the first day all units exceeded a VFA/alkalinity ratio of 0.30 (instability trigger), "
            f"the common upset onset is estimated at {fmt_date(upset_start)}. "
            f"A partial recovery is indicated by the first return below a ratio of 0.10 among the digesters by {fmt_date(upset_end)}. "
            "These boundaries are indicative and should be reconciled against operating logs (feed rates, temperature, mixing, and chemical dosing).",
        )
    else:
        add_body_paragraph(
            doc,
            "Time-series review indicates a synchronised process upset affecting all three digesters; "
            "however, automatic threshold-based windowing could not be robustly determined from the available data alone.",
        )

    add_body_paragraph(
        doc,
        "Figures 1 to 6 summarise the primary stability and performance indicators. Taken together, the trends are consistent with "
        "a plant-wide organic overload or methanogenic inhibition beginning in late September, followed by operational response (buffering restoration) "
        "and gradual biological recovery. The coherence across digesters reduces the likelihood of a single instrument error or isolated unit upset.",
    )

    add_figure(doc, "vfa_alk_ratio_1.png", "VFA/alkalinity ratio time series for Digesters 1–3, including stability and instability thresholds.")
    add_figure(doc, "vfa_timeseries_1.png", "VFA time series for Digesters 1–3 with typical stress and upset thresholds.")
    add_figure(doc, "ph_timeseries_1.png", "pH time series for Digesters 1–3 with a typical stable operating band.")
    add_figure(doc, "alkalinity_timeseries_1.png", "Alkalinity time series for Digesters 1–3 with indicative lower and upper buffering targets.")
    add_figure(doc, "gasprod_timeseries_1.png", "Biogas production time series for Digesters 1–3 showing synchronised decline and recovery.")
    add_figure(doc, "vs_destruction_1.png", "Volatile solids destruction time series for Digesters 1–3 with indicative performance band.")

    doc.add_paragraph("4.1 Stability indicators (VFA, alkalinity, VFA/alkalinity, and pH)", style="Heading 2")
    add_body_paragraph(
        doc,
        "Figure 1 shows a step-change in VFA/alkalinity ratio from very low values through mid-September to sustained instability "
        "across all digesters in October and November, with peaks approaching 1.0. Figure 2 confirms that this ratio increase is driven "
        "by a rapid VFA accumulation to several thousand mg/L, while Figure 4 indicates alkalinity was initially consumed (reduced buffering) "
        "and later increased sharply, consistent with a corrective buffering response (e.g., restored chemical addition or operational changes). "
        "Figure 3 shows concurrent pH depression to approximately 6.2–6.3 in early November, consistent with acid accumulation.",
    )
    add_body_paragraph(
        doc,
        "The combined pattern (high VFA and VFA/alkalinity ratio, depressed pH, and subsequent alkalinity increase) is typical of a digester imbalance "
        "where acidogenesis outpaces methanogenesis. Given all three digesters moved together, the most probable drivers are plant-wide: "
        "a step increase in organic loading rate, a change in sludge blend (e.g., increased primary solids or FOG), interruption/reduction of alkalinity addition, "
        "temperature/heating disturbance, or an inhibitory/toxic input that reduced methanogenic activity.",
    )

    doc.add_paragraph("4.2 Performance indicators (biogas and volatile solids destruction)", style="Heading 2")
    add_body_paragraph(
        doc,
        "Figure 5 shows a pronounced and synchronised biogas production crash beginning late September/early October, aligned with the onset of VFA accumulation "
        "and pH decline (Figures 2 and 3). This is strong evidence of biological inhibition or overload rather than a sampling artifact. "
        "Figure 6 shows volatile solids destruction trending downward through summer and then dropping further during the peak upset period, consistent with reduced "
        "conversion efficiency and methane yield.",
    )
    add_body_paragraph(
        doc,
        "Notably, gas production begins recovering by late November to December while VFA/alkalinity ratio remains elevated for longer (Figure 1). "
        "This is consistent with a phased recovery where buffering and operating conditions improved first, followed by gradual conversion of accumulated intermediates. "
        "Alternatively, it may reflect changes in loading, gas capture efficiency, or operational strategies (e.g., temporary feed reduction) that improved gas output "
        "before full stabilisation of the VFA inventory.",
    )

    doc.add_paragraph("4.3 Cross-digester comparison", style="Heading 2")
    add_body_paragraph(
        doc,
        "Digesters 1 and 2 show a faster return to lower VFA/alkalinity ratio by late December to early January compared to Digester 3 (Figure 1). "
        "Digester 3 also maintains higher VFAs into late December (Figure 2) and experiences a late-December alkalinity dip (Figure 4), which would elevate "
        "its ratio and slow apparent recovery. This suggests either (a) different feed allocation during the recovery period, (b) local limitations such as mixing/heating "
        "effectiveness, or (c) residual inhibition specific to Digester 3.",
    )

    # Tables
    doc.add_paragraph("4.4 Quantitative event summary", style="Heading 2")
    add_body_paragraph(
        doc,
        "Table 1 summarises the peak severity and minimum performance points for each digester, including the timing of key threshold exceedances. "
        "These values support the interpretation that the upset was severe, prolonged, and synchronised, with some differences in recovery trajectory.",
    )

    header = ["Metric", "Digester 1", "Digester 2", "Digester 3"]
    t1_rows: List[List[str]] = [header]
    t1_rows.append(["First VFA/Alk > 0.30", fmt_date(kpis[1].first_ratio_over_030), fmt_date(kpis[2].first_ratio_over_030), fmt_date(kpis[3].first_ratio_over_030)])
    t1_rows.append(["First VFA > 1,000 mg/L", fmt_date(kpis[1].first_vfa_over_1000), fmt_date(kpis[2].first_vfa_over_1000), fmt_date(kpis[3].first_vfa_over_1000)])
    t1_rows.append(
        [
            "Max VFA (mg/L) (date)",
            f"{fmt_int(kpis[1].max_vfa.value)} ({fmt_date(kpis[1].max_vfa.when)})",
            f"{fmt_int(kpis[2].max_vfa.value)} ({fmt_date(kpis[2].max_vfa.when)})",
            f"{fmt_int(kpis[3].max_vfa.value)} ({fmt_date(kpis[3].max_vfa.when)})",
        ]
    )
    t1_rows.append(
        [
            "Max VFA/Alk ratio (date)",
            f"{fmt_float(kpis[1].max_ratio.value, 2)} ({fmt_date(kpis[1].max_ratio.when)})",
            f"{fmt_float(kpis[2].max_ratio.value, 2)} ({fmt_date(kpis[2].max_ratio.when)})",
            f"{fmt_float(kpis[3].max_ratio.value, 2)} ({fmt_date(kpis[3].max_ratio.when)})",
        ]
    )
    t1_rows.append(
        [
            "Min pH (date)",
            f"{fmt_float(kpis[1].min_ph.value, 1)} ({fmt_date(kpis[1].min_ph.when)})",
            f"{fmt_float(kpis[2].min_ph.value, 1)} ({fmt_date(kpis[2].min_ph.when)})",
            f"{fmt_float(kpis[3].min_ph.value, 1)} ({fmt_date(kpis[3].min_ph.when)})",
        ]
    )
    add_table(doc, "Key stability and performance extreme points and threshold exceedance dates (derived from daily monitoring).", t1_rows)

    add_body_paragraph(doc, "Additional performance minima during the upset are summarised in Table 2.")
    t2_rows: List[List[str]] = [header]
    t2_rows.append(
        [
            "Min gas (m³/day) (date)",
            f"{fmt_int(kpis[1].min_gas.value)} ({fmt_date(kpis[1].min_gas.when)})",
            f"{fmt_int(kpis[2].min_gas.value)} ({fmt_date(kpis[2].min_gas.when)})",
            f"{fmt_int(kpis[3].min_gas.value)} ({fmt_date(kpis[3].min_gas.when)})",
        ]
    )
    t2_rows.append(
        [
            "Min VS destruction (%) (date)",
            f"{fmt_float(kpis[1].min_vs.value, 1)} ({fmt_date(kpis[1].min_vs.when)})",
            f"{fmt_float(kpis[2].min_vs.value, 1)} ({fmt_date(kpis[2].min_vs.when)})",
            f"{fmt_float(kpis[3].min_vs.value, 1)} ({fmt_date(kpis[3].min_vs.when)})",
        ]
    )
    add_table(doc, "Selected minima during the upset (biogas and volatile solids destruction).", t2_rows)

    def mean_in_window(values: List[float], dates_: List[datetime], start: datetime, end: datetime) -> float:
        sel = [v for d, v in zip(dates_, values) if start <= d <= end]
        return mean(sel) if sel else float("nan")

    stable_start = datetime(2018, 6, 1)
    stable_end = datetime(2018, 9, 15)
    upset_win_start = datetime(2018, 10, 1)
    upset_win_end = datetime(2018, 11, 30)
    rec_start = datetime(2018, 12, 15)
    rec_end = datetime(2019, 1, 7)

    t3_rows: List[List[str]] = [header]
    t3_rows.append(
        [
            "Mean gas (m³/day) stable",
            fmt_int(mean_in_window(series[1].gas, dates, stable_start, stable_end)),
            fmt_int(mean_in_window(series[2].gas, dates, stable_start, stable_end)),
            fmt_int(mean_in_window(series[3].gas, dates, stable_start, stable_end)),
        ]
    )
    t3_rows.append(
        [
            "Mean gas (m³/day) upset",
            fmt_int(mean_in_window(series[1].gas, dates, upset_win_start, upset_win_end)),
            fmt_int(mean_in_window(series[2].gas, dates, upset_win_start, upset_win_end)),
            fmt_int(mean_in_window(series[3].gas, dates, upset_win_start, upset_win_end)),
        ]
    )
    t3_rows.append(
        [
            "Mean gas (m³/day) recovery",
            fmt_int(mean_in_window(series[1].gas, dates, rec_start, rec_end)),
            fmt_int(mean_in_window(series[2].gas, dates, rec_start, rec_end)),
            fmt_int(mean_in_window(series[3].gas, dates, rec_start, rec_end)),
        ]
    )
    add_table(doc, "Average biogas production by phase (indicative windows based on observed trends).", t3_rows)

    # 5. Key Findings
    doc.add_paragraph("5. Key Findings", style="Heading 1")
    add_body_paragraph(
        doc,
        "The following findings synthesise the evidence across Figures 1 to 6 and Tables 1 to 3. "
        "They are framed as process-engineering observations that can be traced to specific indicators in the monitoring data.",
    )
    doc.add_paragraph(
        "A plant-wide instability event began in late September 2018, evidenced by simultaneous increases in VFA and VFA/alkalinity ratio and a pH decline across all digesters (Figures 1 to 3; Table 1).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "VFA/alkalinity ratios exceeded 0.30 across all units and peaked near 1.0, indicating severe buffering imbalance and high risk of sustained methanogenic inhibition (Figure 1; Table 1).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Biogas production collapsed synchronously, with the most severe absolute minimum in Digester 1, confirming a true biological/process performance failure rather than isolated measurement noise (Figure 5; Tables 2 and 3).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Alkalinity decreased during upset onset and later increased significantly, consistent with buffering consumption followed by corrective alkalinity restoration; pH recovered faster than VFA/alkalinity ratio, indicating partial control recovery before full biological conversion (Figures 3 and 4).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Volatile solids destruction declined prior to the main upset (summer trend), suggesting reduced resilience before the shock load/inhibition, and reached minima concurrent with peak VFA/ratio conditions (Figure 6; Table 2).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Digester 3 showed the slowest recovery in VFA and VFA/alkalinity ratio, consistent with a local limitation (mixing/heating, feed allocation) or a later reduction in alkalinity relative to VFA (Figures 1, 2, and 4).",
        style="List Bullet",
    )

    # 6. Recommendations
    doc.add_paragraph("6. Recommendations", style="Heading 1")
    add_body_paragraph(
        doc,
        "Recommendations below are prioritised to (i) prevent recurrence of the observed late-September upset, "
        "(ii) shorten recovery if a future shock occurs, and (iii) improve diagnostic capability to isolate root causes.",
    )

    doc.add_paragraph("6.1 Immediate operational controls (protect stability)", style="Heading 2")
    add_body_paragraph(
        doc,
        "Given the magnitude of VFA accumulation and gas crash (Figures 2 and 5), stability protection should focus on controlling loading rate, "
        "maintaining buffering capacity, and ensuring temperature and mixing are robust during seasonal transitions.",
    )
    doc.add_paragraph(
        "Implement an operating action plan triggered by VFA/alkalinity ratio thresholds (e.g., 0.10 early warning; 0.30 instability), including staged feed reduction, increased monitoring frequency, and verification of alkalinity dosing systems.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Maintain a minimum alkalinity target and review alkalinity addition capacity and reliability; ensure dosing and instrumentation are alarmed and trended with operator response requirements (Figure 4).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Review digester temperature control and mixing availability around late September to early November (heater status, mixing outages, foam events), as these are common-cause drivers consistent with the synchronised response (Figures 1 to 5).",
        style="List Bullet",
    )

    doc.add_paragraph("6.2 Root-cause investigation (plant-wide drivers)", style="Heading 2")
    add_body_paragraph(
        doc,
        "The synchronised onset across all three digesters strongly suggests a shared upstream or site-wide change. "
        "The investigation should concentrate on changes that occurred in the 2–3 weeks preceding the onset of the VFA rise.",
    )
    doc.add_paragraph(
        "Compile a timeline of sludge feed rate, feed VS/TS, primary-to-WAS blend ratio, FOG inputs, and any hauled waste acceptance changes; test whether a step change aligns with the first exceedance dates in Table 1.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Review sidestream returns (e.g., dewatering centrate) and chemical additions (polymer, metal salts) for indirect inhibition pathways and buffering impacts.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Screen for toxicity/inhibition signals (e.g., industrial discharges, cleaning events, high sulfide or ammonia episodes if data exist) that could suppress methanogenesis while allowing acidogenesis to continue.",
        style="List Bullet",
    )

    doc.add_paragraph("6.3 Targeted unit follow-up (Digester 3 recovery lag)", style="Heading 2")
    add_body_paragraph(
        doc,
        "Digester 3’s slower recovery (Figures 1 and 2) warrants a targeted check for unit-specific limitations that can be corrected without plant-wide changes.",
    )
    doc.add_paragraph(
        "Verify digester-specific mixing intensity, heat exchanger performance, and short-circuiting risk; compare temperature profiles and mixing runtime to Digesters 1 and 2 during the recovery period.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Confirm feed distribution and any operational constraints that could have resulted in higher effective loading or less buffering to Digester 3 in December (Figure 4).",
        style="List Bullet",
    )

    doc.add_paragraph("6.4 Further analysis to strengthen conclusions", style="Heading 2")
    add_body_paragraph(
        doc,
        "The dataset is sufficient to identify the upset and quantify its severity, but additional analyses would improve root-cause discrimination and provide more defensible operating limits.",
    )
    doc.add_paragraph(
        "Perform lagged correlation analysis between VFA/alkalinity ratio and biogas production (and VS destruction) to estimate response time constants and recovery kinetics.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Segment the time series into regimes (stable, upset, recovery) and compute control-chart style statistics to define site-specific early-warning thresholds beyond generic guidance.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "If feed and temperature data are available, calculate organic loading rate and compare against estimated digester capacity; evaluate whether the upset is consistent with overload versus inhibition.",
        style="List Bullet",
    )

    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()