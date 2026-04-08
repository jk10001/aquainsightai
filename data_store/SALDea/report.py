# filename: report.py
"""
DOCX report generator (revised per user instructions)

Checklist (before generating the report):
- Verify expected figure/CSV inputs exist; gracefully fall back to alternate filenames where provided.
- Rebuild body tables to show pooled-only results, and move full by-plant tables to an appendix with clear cross-references.
- Combine specified table pairs (influent, effluent, and LRV summaries) and apply rounding/thousands separators and reduced table font where requested.
- Ensure all figures/tables have automatic-numbered captions and are referenced in the narrative by table/figure number.
- Apply consistent A4 layout, header/footer styling (10 pt grey), and save as report.docx in the current directory.
"""

from __future__ import annotations

import csv
import math
import os
from datetime import date
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# ────────────────────────────────────────────────────────────────
#                       low-level helpers
# ────────────────────────────────────────────────────────────────
def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Optional[Pt] = None,
    font_color: Optional[RGBColor] = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(doc: Document, title: str, report_date: Optional[str] = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    # Ensure footer style (incl. fields) is 10pt grey where available
    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style
        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: Optional[str] = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    doc.add_paragraph(title, style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(report_date, style="Subtitle").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


# ────────────────────────────────────────────────────────────────
#                       formatting helpers
# ────────────────────────────────────────────────────────────────
def _safe_float(x: object) -> Optional[float]:
    if x is None:
        return None
    if isinstance(x, str):
        s = x.strip()
        if s == "" or s.lower() in {"nan", "n/a"}:
            return None
        try:
            return float(s)
        except Exception:
            return None
    try:
        v = float(x)
    except Exception:
        return None
    if math.isnan(v):
        return None
    return v


def format_num(x: object, decimals: int = 2, use_commas: bool = True) -> str:
    """
    General numeric formatter with optional thousands separators.
    """
    if x is None:
        return "N/A"
    if isinstance(x, str):
        s = x.strip()
        if s == "" or s.lower() in {"nan", "n/a"}:
            return "N/A"
        return s

    v = _safe_float(x)
    if v is None:
        return str(x)

    if decimals < 0:
        decimals = 0

    if use_commas:
        s = f"{v:,.{decimals}f}"
    else:
        s = f"{v:.{decimals}f}"

    if "." in s and decimals > 0:
        s = s.rstrip("0").rstrip(".")
    return s


def format_pct(x: object, decimals: int = 1) -> str:
    v = _safe_float(x)
    if v is None:
        return "N/A"
    return f"{v:.{decimals}f}"


def read_csv_as_dicts(path: str) -> List[Dict[str, str]]:
    with open(path, "r", newline="", encoding="utf-8") as f:
        rdr = csv.DictReader(f)
        return [dict(row) for row in rdr]


def parse_inline_csv(text: str) -> List[Dict[str, str]]:
    lines = [ln for ln in text.strip().splitlines() if ln.strip()]
    rdr = csv.DictReader(lines)
    return [dict(row) for row in rdr]


def ensure_files_exist(paths: List[str]) -> List[str]:
    return [p for p in paths if not os.path.exists(p)]


def pick_first_existing(candidates: List[str]) -> Optional[str]:
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def add_paragraph_justified(doc: Document, text: str, style: str = "Normal") -> Paragraph:
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p


def set_table_font(table, size_pt: int = 9, color: Optional[RGBColor] = None) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)
                    if color is not None:
                        run.font.color.rgb = color


def add_table_generic(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
    font_size_pt: Optional[int] = None,
    total_width_mm: float = 160.0,
) -> None:
    if not rows:
        return

    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.rows[0].repeat_header = True

    # Set widths to fit page content region
    num_cols = len(rows[0])
    if num_cols > 0:
        col_w = total_width_mm / num_cols
        for col in t.columns:
            col.width = Mm(col_w)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    if font_size_pt is not None:
        set_table_font(t, size_pt=font_size_pt)

    doc.add_paragraph()


def add_figure(doc: Document, image_path: str, caption: str, width_mm: float = 160.0) -> None:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Mm(width_mm))
        add_caption(doc, "Figure", caption)
        doc.add_paragraph()
    else:
        add_paragraph_justified(doc, f"Figure file not found: {image_path}")


# ────────────────────────────────────────────────────────────────
#                       content mapping helpers
# ────────────────────────────────────────────────────────────────
def plant_full_name(short: str) -> str:
    mapping = {"NoWe": "Northwest El Paso", "BiSp": "Big Spring", "BiSP": "Big Spring", "BrWo": "Brownwood"}
    return mapping.get(short, short)


# ────────────────────────────────────────────────────────────────
#                       table builders
# ────────────────────────────────────────────────────────────────
def build_prevalence_table(rows: List[Dict[str, str]], stream: str, pooled_only: bool) -> List[List[str]]:
    """
    stream: 'Influent' or 'Effluent'
    pooled_only: if True, returns only "All plants pooled" row(s) and excludes Plant col.
    """
    if stream.lower() == "influent":
        n_key = "n_influent_samples [-]"
        det_key = "n_influent_detected [-]"
        prev_key = "influent_detection_prevalence [%]"
    else:
        n_key = "n_effluent_samples [-]"
        det_key = "n_effluent_detected [-]"
        prev_key = "effluent_detection_prevalence [%]"

    if pooled_only:
        out = [["Microorganism", "n", "Detections", "Prevalence (%)"]]
    else:
        out = [["Plant", "Microorganism", "n", "Detections", "Prevalence (%)"]]

    for r in rows:
        if pooled_only and r.get("Plant") != "All plants pooled":
            continue
        if (not pooled_only) and r.get("Plant") == "All plants pooled":
            # keep pooled row for full table as well
            pass

        if pooled_only:
            out.append(
                [
                    r.get("Microorganism", ""),
                    format_num(r.get(n_key), decimals=0),
                    format_num(r.get(det_key), decimals=0),
                    format_pct(r.get(prev_key), decimals=1),
                ]
            )
        else:
            out.append(
                [
                    r.get("Plant", ""),
                    r.get("Microorganism", ""),
                    format_num(r.get(n_key), decimals=0),
                    format_num(r.get(det_key), decimals=0),
                    format_pct(r.get(prev_key), decimals=1),
                ]
            )

    # Sort with pooled last for full tables; alphabetical micro for pooled-only
    header = out[0]
    data = out[1:]
    if pooled_only:
        data = sorted(data, key=lambda rr: rr[0])
    else:
        def key_row(rr: List[str]) -> Tuple[int, str, str]:
            return (1 if rr[0] == "All plants pooled" else 0, rr[0], rr[1])
        data = sorted(data, key=key_row)

    return [header] + data


def build_stream_stats_combined(
    stats_rows: List[Dict[str, str]],
    stream: str,
    pooled_only: bool,
    decimals: int,
) -> List[List[str]]:
    """
    Combined table (median, mean, p25, p75) for Influent or Effluent.
    If pooled_only, excludes Plant column.
    """
    stream_l = stream.strip().lower()
    filt = [r for r in stats_rows if r.get("Stream", "").strip().lower() == stream_l]
    if pooled_only:
        filt = [r for r in filt if r.get("Plant") == "All plants pooled"]

    if pooled_only:
        out = [["Microorganism", "n", "Median", "Mean", "P25", "P75"]]
    else:
        out = [["Plant", "Microorganism", "n", "Median", "Mean", "P25", "P75"]]

    for r in filt:
        if pooled_only:
            out.append(
                [
                    r.get("Microorganism", ""),
                    format_num(r.get("n_samples [-]"), decimals=0),
                    format_num(_safe_float(r.get("median [original units]")), decimals=decimals),
                    format_num(_safe_float(r.get("mean [original units]")), decimals=decimals),
                    format_num(_safe_float(r.get("p25 [original units]")), decimals=decimals),
                    format_num(_safe_float(r.get("p75 [original units]")), decimals=decimals),
                ]
            )
        else:
            out.append(
                [
                    r.get("Plant", ""),
                    r.get("Microorganism", ""),
                    format_num(r.get("n_samples [-]"), decimals=0),
                    format_num(_safe_float(r.get("median [original units]")), decimals=decimals),
                    format_num(_safe_float(r.get("mean [original units]")), decimals=decimals),
                    format_num(_safe_float(r.get("p25 [original units]")), decimals=decimals),
                    format_num(_safe_float(r.get("p75 [original units]")), decimals=decimals),
                ]
            )

    header = out[0]
    data = out[1:]
    if pooled_only:
        data = sorted(data, key=lambda rr: rr[0])
    else:
        def key_row(rr: List[str]) -> Tuple[int, str, str]:
            return (1 if rr[0] == "All plants pooled" else 0, rr[0], rr[1])
        data = sorted(data, key=key_row)

    return [header] + data


def build_lrv_combined(
    lrv_rows: List[Dict[str, str]],
    pooled_only: bool,
    decimals: int,
) -> List[List[str]]:
    """
    Combined LRV table (median, mean, p25, p75).
    If pooled_only, excludes Plant column.
    """
    if pooled_only:
        filt = [r for r in lrv_rows if r.get("Plant") == "All plants pooled"]
        out = [["Microorganism", "n pairs", "Median LRV (log10)", "Mean LRV (log10)", "P25 LRV (log10)", "P75 LRV (log10)"]]
    else:
        filt = list(lrv_rows)
        out = [["Plant", "Microorganism", "n pairs", "Median LRV (log10)", "Mean LRV (log10)", "P25 LRV (log10)", "P75 LRV (log10)"]]

    for r in filt:
        med = _safe_float(r.get("median_LRV [log10]"))
        mean = _safe_float(r.get("mean_LRV [log10]"))
        p25 = _safe_float(r.get("p25_LRV [log10]"))
        p75 = _safe_float(r.get("p75_LRV [log10]"))

        if pooled_only:
            out.append(
                [
                    r.get("Microorganism", ""),
                    format_num(r.get("n_pairs [-]"), decimals=0),
                    format_num(med, decimals=decimals, use_commas=False),
                    format_num(mean, decimals=decimals, use_commas=False),
                    format_num(p25, decimals=decimals, use_commas=False),
                    format_num(p75, decimals=decimals, use_commas=False),
                ]
            )
        else:
            out.append(
                [
                    r.get("Plant", ""),
                    r.get("Microorganism", ""),
                    format_num(r.get("n_pairs [-]"), decimals=0),
                    format_num(med, decimals=decimals, use_commas=False),
                    format_num(mean, decimals=decimals, use_commas=False),
                    format_num(p25, decimals=decimals, use_commas=False),
                    format_num(p75, decimals=decimals, use_commas=False),
                ]
            )

    header = out[0]
    data = out[1:]
    if pooled_only:
        data = sorted(data, key=lambda rr: rr[0])
    else:
        def key_row(rr: List[str]) -> Tuple[int, str, str]:
            return (1 if rr[0] == "All plants pooled" else 0, rr[0], rr[1])
        data = sorted(data, key=key_row)

    return [header] + data


def build_correlation_table(corr_rows: List[Dict[str, str]]) -> List[List[str]]:
    rows = [["Pathogen", "Indicator", "n pairs", "Spearman ρ", "Pearson r"]]
    for r in corr_rows:
        rows.append(
            [
                r.get("Pathogen", ""),
                r.get("Indicator", ""),
                format_num(r.get("n_pairs [-]"), decimals=0),
                format_num(_safe_float(r.get("Spearman_rho [-]")), decimals=3, use_commas=False),
                format_num(_safe_float(r.get("Pearson_r [-]")), decimals=3, use_commas=False),
            ]
        )
    rows_data = sorted(rows[1:], key=lambda rr: (rr[0], rr[1]))
    return [rows[0]] + rows_data


# ────────────────────────────────────────────────────────────────
#                       MAIN report generator
# ────────────────────────────────────────────────────────────────
def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Wastewater Microbial Indicators and Pathogens: Prevalence, Removal, and Surrogate Assessment"
    today_str = date.today().strftime("%d %B %Y")

    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(doc, include_toc=True, include_lof=True, include_lot=True)

    # Inline CSV for prevalence table (provided in prompt)
    prevalence_inline_csv = """
Plant,Microorganism,Units,n_influent_samples [-],n_influent_detected [-],influent_detection_prevalence [%],n_effluent_samples [-],n_effluent_detected [-],effluent_detection_prevalence [%]
BiSp,Adenovirus,MPN/L,16,16,100.0,16,13,81.2
BiSp,Aerobic endospores,CFU/100mL,14,14,100.0,14,14,100.0
BiSp,Cryptosporidium,/L,16,11,68.8,16,12,75.0
BiSp,E. coli,MPN/100mL,14,14,100.0,13,8,61.5
BiSp,Fecal coliform,MPN/100mL,14,14,100.0,13,10,76.9
BiSp,Giardia,/L,16,16,100.0,16,14,87.5
BiSp,Male-specific coliphage,PFU/mL,14,14,100.0,15,1,6.7
BiSp,Somatic coliphage,PFU/mL,13,13,100.0,15,2,13.3
BiSp,Total coliform,MPN/100mL,13,13,100.0,12,11,91.7
BrWo,Adenovirus,MPN/L,16,14,87.5,16,6,37.5
BrWo,Aerobic endospores,CFU/100mL,15,15,100.0,15,15,100.0
BrWo,Cryptosporidium,/L,15,10,66.7,15,14,93.3
BrWo,E. coli,MPN/100mL,15,15,100.0,15,7,46.7
BrWo,Fecal coliform,MPN/100mL,15,15,100.0,15,7,46.7
BrWo,Giardia,/L,15,14,93.3,15,15,100.0
BrWo,Male-specific coliphage,PFU/mL,15,15,100.0,15,1,6.7
BrWo,Somatic coliphage,PFU/mL,15,15,100.0,15,2,13.3
BrWo,Total coliform,MPN/100mL,14,14,100.0,14,13,92.9
NoWe,Adenovirus,MPN/L,16,16,100.0,16,9,56.2
NoWe,Aerobic endospores,CFU/100mL,16,16,100.0,16,16,100.0
NoWe,Cryptosporidium,/L,16,10,62.5,16,13,81.2
NoWe,E. coli,MPN/100mL,16,16,100.0,16,16,100.0
NoWe,Fecal coliform,MPN/100mL,16,16,100.0,16,16,100.0
NoWe,Giardia,/L,16,16,100.0,16,16,100.0
NoWe,Male-specific coliphage,PFU/mL,16,16,100.0,16,3,18.8
NoWe,Somatic coliphage,PFU/mL,16,16,100.0,16,5,31.2
NoWe,Total coliform,MPN/100mL,15,15,100.0,15,15,100.0
All plants pooled,Adenovirus,MPN/L,48,46,95.8,48,28,58.3
All plants pooled,Aerobic endospores,CFU/100mL,45,45,100.0,45,45,100.0
All plants pooled,Cryptosporidium,/L,47,31,66.0,47,39,83.0
All plants pooled,E. coli,MPN/100mL,45,45,100.0,44,31,70.5
All plants pooled,Fecal coliform,MPN/100mL,45,45,100.0,44,33,75.0
All plants pooled,Giardia,/L,47,46,97.9,47,45,95.7
All plants pooled,Male-specific coliphage,PFU/mL,45,45,100.0,46,5,10.9
All plants pooled,Somatic coliphage,PFU/mL,44,44,100.0,46,9,19.6
All plants pooled,Total coliform,MPN/100mL,42,42,100.0,41,39,95.1
""".strip()
    prevalence_rows = parse_inline_csv(prevalence_inline_csv)

    # Read remaining analysis outputs (if present)
    stats_path = "summary_statistics_influent_effluent_by_plant_pooled.csv"
    lrv_path = "log_removal_value_summary_by_plant_pooled.csv"
    corr_path = "correlation_indicator_pathogen_lrv.csv"
    stats_rows = read_csv_as_dicts(stats_path) if os.path.exists(stats_path) else []
    lrv_rows = read_csv_as_dicts(lrv_path) if os.path.exists(lrv_path) else []
    corr_rows = read_csv_as_dicts(corr_path) if os.path.exists(corr_path) else []

    # Figures (use revised filenames where provided)
    fig_indicator_by_plant = pick_first_existing(
        ["lrv_boxplots_by_plant_revision_v2.png", "lrv_boxplots_by_plant.png"]
    )
    fig_protozoa_by_plant = pick_first_existing(["lrv_boxplots_protozoa_by_plant.png"])
    fig_adeno_by_plant = pick_first_existing(["lrv_boxplot_adenovirus_by_plant.png"])
    fig_pooled_all = pick_first_existing(
        ["lrv_boxplot_all_microorganisms_pooled_no_markers.png", "lrv_boxplot_all_microorganisms.png"]
    )
    fig_scatter_adeno = pick_first_existing(["scatter_adenovirus_vs_indicators.png"])
    fig_scatter_giardia = pick_first_existing(["scatter_giardia_vs_indicators.png"])
    fig_scatter_crypto = pick_first_existing(["scatter_cryptosporidium_vs_indicators.png"])

    expected_files = [
        p for p in [
            fig_indicator_by_plant,
            fig_protozoa_by_plant,
            fig_adeno_by_plant,
            fig_pooled_all,
            fig_scatter_adeno,
            fig_scatter_giardia,
            fig_scatter_crypto,
            stats_path,
            lrv_path,
            corr_path,
        ] if p is not None
    ]
    missing = ensure_files_exist(expected_files)

    # Plant descriptions (include short names in Table 1 per instruction)
    plants_info = [
        ("NoWe", "Northwest El Paso",
         "Primary clarification, activated sludge, secondary clarification, disc filter microscreens", "UV"),
        ("BiSp", "Big Spring",
         "Primary clarification, activated sludge/trickling filter, secondary clarification", "Chlorine"),
        ("BrWo", "Brownwood",
         "Primary clarification, activated sludge, secondary clarification", "Chlorine"),
    ]

    # ───────────────────────────────────────────────────────────
    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    add_paragraph_justified(
        doc,
        "This report summarizes indicator microorganisms and pathogenic targets measured across three wastewater treatment plants "
        "between March 2015 and June 2016. The dataset includes bacterial indicators (total coliform, fecal coliform, E. coli), "
        "bacteriophages (somatic and male-specific coliphages), aerobic endospores, protozoan parasites (Giardia and Cryptosporidium), "
        "and adenovirus. Results are interpreted in terms of detection prevalence and log-removal values (LRVs) to support treatment "
        "performance understanding and surrogate selection for pathogen control.",
    )

    # ───────────────────────────────────────────────────────────
    # 2. Objectives
    doc.add_heading("2. Objectives", level=1)
    add_paragraph_justified(
        doc,
        "The assessment focuses on three questions that directly support process performance evaluation and monitoring strategy development:",
    )
    doc.add_paragraph(
        "Quantify prevalence of indicator and pathogen detections in influent and effluent across plants.",
        style="List Number",
    )
    doc.add_paragraph(
        "Evaluate microbial removal through treatment (LRVs), including plant-to-plant comparisons.",
        style="List Number",
    )
    doc.add_paragraph(
        "Assess the extent to which indicator LRVs track pathogen LRVs (surrogate effectiveness) when plants are pooled.",
        style="List Number",
    )
    doc.add_paragraph()

    # ───────────────────────────────────────────────────────────
    # 3. Description of Dataset
    doc.add_heading("3. Description of Dataset", level=1)
    add_paragraph_justified(
        doc,
        "The source workbook (WWTP_microbial_loads_and_removal.xlsx) contains event-based measurements by plant, sampling month, "
        "sample type (influent grab, effluent grab, UF effluent), and microorganism target. Non-detects are represented as BDL "
        "or numeric zero; these were treated as method detection limit (MDL) values in the underlying analysis. For protozoa and "
        "coliphages, the UF effluent sample was used as the effluent representation; for all other targets, the effluent grab sample was used.",
    )
    add_paragraph_justified(
        doc,
        "The three plants and their major treatment and disinfection configurations are summarized below to support interpretation of plant-to-plant differences:",
    )
    plant_tbl_rows = [["Plant (short name)", "Plant (full name)", "Primary/secondary treatment train", "Disinfection"]]
    for short, full, train, disinf in plants_info:
        plant_tbl_rows.append([short, full, train, disinf])
    add_table_generic(
        doc,
        "Wastewater treatment plants included in the dataset and their disinfection configurations.",
        plant_tbl_rows,
        total_width_mm=160.0,
        font_size_pt=9,
    )
    add_paragraph_justified(
        doc,
        "Plant short names used throughout are: NoWe (Northwest El Paso), BiSp (Big Spring), and BrWo (Brownwood).",
    )

    if missing:
        add_paragraph_justified(
            doc,
            "Some expected input files were not found in the current directory and could not be embedded in the report: "
            + ", ".join(missing),
        )

    # ───────────────────────────────────────────────────────────
    # 4. Analysis
    doc.add_heading("4. Analysis", level=1)

    # 4.1 Detection prevalence
    doc.add_heading("4.1 Detection prevalence (influent vs effluent)", level=2)
    add_paragraph_justified(
        doc,
        "Detection prevalence provides a first-pass view of how consistently each target occurs and whether effluent results are frequently "
        "censored by the MDL. Tables 2 and 3 summarize pooled detection prevalence across all plants. Full by-plant prevalence tables are "
        "provided in the Appendix (Tables 8 and 9).",
    )

    infl_prev_pooled = build_prevalence_table(prevalence_rows, stream="Influent", pooled_only=True)
    eff_prev_pooled = build_prevalence_table(prevalence_rows, stream="Effluent", pooled_only=True)

    add_table_generic(
        doc,
        "Influent detection prevalence (all plants pooled).",
        infl_prev_pooled,
        total_width_mm=150.0,
        font_size_pt=9,
    )
    add_table_generic(
        doc,
        "Effluent detection prevalence (all plants pooled).",
        eff_prev_pooled,
        total_width_mm=150.0,
        font_size_pt=9,
    )

    add_paragraph_justified(
        doc,
        "Across all plants, influent detections are near-universal for bacterial indicators and both coliphage groups, whereas influent "
        "Cryptosporidium prevalence is materially lower than other targets. In effluent, coliphages show frequent non-detects while aerobic "
        "endospores remain consistently detected, indicating persistence through the treatment train and limited sensitivity to disinfection.",
    )

    # 4.2 Influent and effluent concentration statistics
    doc.add_heading("4.2 Influent and effluent concentration statistics", level=2)
    add_paragraph_justified(
        doc,
        "Concentration summaries in original units provide context for interpreting LRVs, particularly where effluent values are censored at "
        "the MDL. Tables 4 and 5 present pooled influent and effluent statistics. Full by-plant concentration statistics are provided in the "
        "Appendix (Tables 10 and 11).",
    )

    if stats_rows:
        # Per instruction: combine central tendency + percentiles, round to 0 decimals, thousands separators, reduced font in full tables.
        # Body: pooled only, no Plant column.
        infl_stats_pooled = build_stream_stats_combined(stats_rows, stream="Influent", pooled_only=True, decimals=0)
        eff_stats_pooled = build_stream_stats_combined(stats_rows, stream="Effluent", pooled_only=True, decimals=0)

        add_table_generic(
            doc,
            "Influent concentrations summary (all plants pooled; original units).",
            infl_stats_pooled,
            total_width_mm=160.0,
            font_size_pt=8,
        )
        add_table_generic(
            doc,
            "Effluent concentrations summary (all plants pooled; original units).",
            eff_stats_pooled,
            total_width_mm=160.0,
            font_size_pt=8,
        )
    else:
        add_paragraph_justified(doc, "Influent/effluent summary-statistics CSV was not available for inclusion.")

    add_paragraph_justified(
        doc,
        "Pooled results indicate that bacterial indicators occur at very high influent levels (orders of magnitude higher than protozoa and adenovirus). "
        "Effluent distributions for coliphages are strongly influenced by MDL censoring (median effluent values at or near the MDL), whereas adenovirus "
        "shows evidence of episodic breakthrough in some events (reflected by a high effluent mean relative to the median). These patterns are consistent "
        "with strong baseline treatment performance punctuated by event-driven variability that is more apparent for pathogens than for bacterial indicators.",
    )

    # 4.3 LRV distributions by plant
    doc.add_heading("4.3 Log-removal values by plant", level=2)
    add_paragraph_justified(
        doc,
        "Figures 1 to 3 compare LRV distributions by plant, separated into indicator organisms, protozoan parasites, and adenovirus. Markers differentiate "
        "events where the effluent was non-detect (MDL-censored) versus detected, which is important for interpreting upper tails that may be driven by substitution "
        "at low effluent values.",
    )

    if fig_indicator_by_plant:
        add_figure(
            doc,
            fig_indicator_by_plant,
            "Indicator microorganism LRVs by plant (non-detect and detect effluent markers).",
            width_mm=175.0,
        )
    if fig_protozoa_by_plant:
        add_figure(
            doc,
            fig_protozoa_by_plant,
            "Protozoan parasite LRVs by plant (non-detect and detect effluent markers).",
            width_mm=175.0,
        )
    if fig_adeno_by_plant:
        add_figure(
            doc,
            fig_adeno_by_plant,
            "Adenovirus LRVs by plant (non-detect and detect effluent markers).",
            width_mm=155.0,
        )

    add_paragraph_justified(
        doc,
        "Bacterial indicators exhibit consistently high LRVs (typically around 5 to 7 log10), indicating strong reduction through the combined primary and "
        "secondary trains and disinfection. In contrast, protozoa show substantially lower and more variable removal, consistent with greater reliance on physical "
        "separation and solids/particle capture rather than disinfection. Adenovirus exhibits wide within-plant spread, consistent with disinfection sensitivity and "
        "episodic performance excursions.",
    )

    # 4.4 Pooled LRV patterns and summary statistics
    doc.add_heading("4.4 Pooled LRV patterns and summary statistics", level=2)
    add_paragraph_justified(
        doc,
        "Figure 4 shows LRVs pooled across plants, highlighting distinct LRV bands by microorganism group. Table 6 summarizes pooled LRV statistics "
        "(median, mean, and interquartile range). Full by-plant LRV statistics are provided in the Appendix (Table 12).",
    )
    if fig_pooled_all:
        add_figure(
            doc,
            fig_pooled_all,
            "LRV distribution pooled across plants for all microorganisms.",
            width_mm=175.0,
        )

    if lrv_rows:
        lrv_pooled = build_lrv_combined(lrv_rows, pooled_only=True, decimals=2)
        add_table_generic(
            doc,
            "Log-removal value (LRV) summary statistics (all plants pooled; log10).",
            lrv_pooled,
            total_width_mm=160.0,
            font_size_pt=9,
        )
    else:
        add_paragraph_justified(doc, "LRV summary-statistics CSV was not available for inclusion.")

    add_paragraph_justified(
        doc,
        "Pooled LRV patterns indicate that aerobic endospores have the lowest removal among the indicators, consistent with persistence and association with "
        "solids. Coliphages occupy an intermediate removal band. Bacterial indicators show the highest LRVs and the most compressed spread, which limits their "
        "ability to resolve pathogen removal variability when data are pooled across plants and events.",
    )

    # 4.5 Surrogate effectiveness
    doc.add_heading("4.5 Surrogate effectiveness: pathogen LRV vs indicator LRV", level=2)
    add_paragraph_justified(
        doc,
        "Surrogacy was evaluated using event-matched comparisons of pathogen LRVs versus indicator LRVs with all plants pooled. Figures 5 to 7 show scatter plots "
        "for adenovirus, Giardia, and Cryptosporidium, respectively. Table 7 summarizes Spearman and Pearson correlations for the same matched datasets.",
    )

    if fig_scatter_adeno:
        add_figure(
            doc,
            fig_scatter_adeno,
            "Event-matched scatter: adenovirus LRV versus indicator LRVs (all plants pooled).",
            width_mm=175.0,
        )
    if fig_scatter_giardia:
        add_figure(
            doc,
            fig_scatter_giardia,
            "Event-matched scatter: Giardia LRV versus indicator LRVs (all plants pooled).",
            width_mm=175.0,
        )
    if fig_scatter_crypto:
        add_figure(
            doc,
            fig_scatter_crypto,
            "Event-matched scatter: Cryptosporidium LRV versus indicator LRVs (all plants pooled).",
            width_mm=175.0,
        )

    if corr_rows:
        corr_tbl = build_correlation_table(corr_rows)
        add_table_generic(
            doc,
            "Correlation between indicator LRVs and pathogen LRVs (all plants pooled; matched events).",
            corr_tbl,
            total_width_mm=160.0,
            font_size_pt=9,
        )
    else:
        add_paragraph_justified(doc, "Correlation CSV was not available for inclusion.")

    add_paragraph_justified(
        doc,
        "When pooled across plants, surrogate relationships are generally weak. Bacterial indicator LRVs are consistently high and compressed, producing vertical "
        "point clouds against pathogen LRVs that vary more widely. Coliphages provide more spread in indicator LRV but still show substantial pathogen variability at "
        "similar coliphage LRV. These results support using indicators primarily for operational verification within a defined treatment context, rather than as "
        "stand-alone predictors of virus or protozoa removal across different plants and configurations.",
    )

    # ───────────────────────────────────────────────────────────
    # 5. Key Findings
    doc.add_heading("5. Key Findings", level=1)
    add_paragraph_justified(
        doc,
        "Key findings are summarized below, integrating pooled detection prevalence (Tables 2 and 3), pooled concentration context (Tables 4 and 5), "
        "removal performance (Figure 4 and Table 6), and surrogate analyses (Figures 5 to 7 and Table 7).",
    )
    doc.add_paragraph(
        "Influent detections were near-universal for bacterial indicators and coliphages, while Cryptosporidium influent prevalence was materially lower than other targets (Table 2).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Effluent detections differed strongly by organism group: aerobic endospores remained detected in essentially all effluents, while coliphages were frequently non-detect, indicating strong reduction and/or MDL censoring (Table 3).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Bacterial indicators achieved the highest LRVs overall (typically around 5 to 7 log10), demonstrating robust removal of bacterial targets, but with limited ability to resolve event-to-event pathogen variability (Figure 4 and Table 6).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Protozoan removal (Giardia and Cryptosporidium) was substantially lower than bacterial indicators, indicating that fecal indicator LRVs overstate protozoa removal performance when used as surrogates (Figure 4 and Table 6).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Adenovirus removal showed high variability relative to indicators, consistent with disinfection sensitivity and episodic performance excursions (Figure 4).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Correlations between indicator LRVs and pathogen LRVs were generally weak when pooled across plants, indicating limited surrogate performance without additional process context (Table 7).",
        style="List Bullet",
    )
    doc.add_paragraph()

    # ───────────────────────────────────────────────────────────
    # 6. Recommendations
    doc.add_heading("6. Recommendations", level=1)
    add_paragraph_justified(
        doc,
        "Recommendations are provided to strengthen performance assessment and improve the practical use of indicators for operational control "
        "and risk-based decision-making across the three plants.",
    )
    doc.add_paragraph(
        "Maintain direct pathogen monitoring for protozoa and viruses where risk targets depend on these organisms; bacterial fecal indicators alone are not sufficient to infer protozoa or adenovirus removal across different treatment trains (Figures 5 to 7 and Table 7).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Use coliphages as a potential intermediate indicator where viral surrogacy is desired, but validate performance relationships within each plant and treatment configuration rather than relying on pooled relationships (Table 7).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "For plants with disinfection-driven variability in virus removal, prioritize verification of disinfection performance during low-LRV events (e.g., UV dose delivery and UV transmittance for UV systems; disinfectant residual and contact time for chlorine systems), and link microbial excursions to operational records where possible.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "For protozoa control, focus on unit processes responsible for physical removal and particle capture (e.g., clarification, filtration/microscreens). Where available, add complementary surrogates tied to physical removal (turbidity, particle counts, TSS, filter headloss) to support mechanistic interpretation of protozoa LRV variability.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Further analyses recommended: (1) plant-specific surrogate assessments (correlations and conditional performance) rather than pooled; (2) seasonal/monthly stratification to identify weather- or temperature-linked performance changes; (3) censoring-aware statistics for MDL-heavy targets such as coliphages; and (4) event-based investigations linking low LRVs to operational changes, hydraulics, and unit process performance.",
        style="List Bullet",
    )
    doc.add_paragraph()

    # ───────────────────────────────────────────────────────────
    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix A. Full Tables by Plant", level=1)

    add_paragraph_justified(
        doc,
        "This appendix provides the full by-plant tables that underpin the pooled summaries presented in the main body of the report.",
    )

    # Appendix: full prevalence tables
    doc.add_heading("Appendix A.1 Detection prevalence (by plant and pooled)", level=2)
    add_paragraph_justified(
        doc,
        "Tables 8 and 9 provide influent and effluent detection prevalence by plant and pooled across all plants.",
    )
    infl_prev_full = build_prevalence_table(prevalence_rows, stream="Influent", pooled_only=False)
    eff_prev_full = build_prevalence_table(prevalence_rows, stream="Effluent", pooled_only=False)
    add_table_generic(
        doc,
        "Influent detection prevalence by plant and pooled (all targets).",
        infl_prev_full,
        total_width_mm=160.0,
        font_size_pt=8,
    )
    add_table_generic(
        doc,
        "Effluent detection prevalence by plant and pooled (all targets).",
        eff_prev_full,
        total_width_mm=160.0,
        font_size_pt=8,
    )

    # Appendix: full concentration tables (combined Table 4+5 and 6+7 equivalents)
    doc.add_heading("Appendix A.2 Concentration statistics (by plant and pooled)", level=2)
    add_paragraph_justified(
        doc,
        "Tables 10 and 11 provide combined concentration statistics (median, mean, and interquartile range) for influent and effluent, by plant and pooled.",
    )
    if stats_rows:
        infl_stats_full = build_stream_stats_combined(stats_rows, stream="Influent", pooled_only=False, decimals=0)
        eff_stats_full = build_stream_stats_combined(stats_rows, stream="Effluent", pooled_only=False, decimals=0)

        add_table_generic(
            doc,
            "Influent concentrations summary by plant and pooled (original units).",
            infl_stats_full,
            total_width_mm=180.0,
            font_size_pt=7,
        )
        add_table_generic(
            doc,
            "Effluent concentrations summary by plant and pooled (original units).",
            eff_stats_full,
            total_width_mm=180.0,
            font_size_pt=7,
        )
    else:
        add_paragraph_justified(doc, "Influent/effluent summary-statistics CSV was not available for inclusion.")

    # Appendix: full LRV table (combined Table 8+9 equivalent)
    doc.add_heading("Appendix A.3 LRV statistics (by plant and pooled)", level=2)
    add_paragraph_justified(
        doc,
        "Table 12 provides combined LRV summary statistics (median, mean, and interquartile range) by plant and pooled across all plants.",
    )
    if lrv_rows:
        lrv_full = build_lrv_combined(lrv_rows, pooled_only=False, decimals=2)
        add_table_generic(
            doc,
            "Log-removal value (LRV) summary statistics by plant and pooled (log10).",
            lrv_full,
            total_width_mm=190.0,
            font_size_pt=7,
        )
    else:
        add_paragraph_justified(doc, "LRV summary-statistics CSV was not available for inclusion.")

    # Save
    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()