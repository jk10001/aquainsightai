# filename: report.py
from datetime import date
from typing import List
import csv
from io import StringIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


def configure_page(doc: Document) -> None:
    for s in doc.sections:
        s.page_height = Mm(297)
        s.page_width = Mm(210)
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(
    doc: Document, title: str, report_date: str | None = None
) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        footer_style = None

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style
        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style
        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    title_p = doc.add_paragraph(title, style="Title")
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_p = doc.add_paragraph(report_date, style="Subtitle")
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    if not rows:
        return
    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    t.rows[0].repeat_header = True

    available_width_mm = 160
    num_cols = len(rows[0])
    if num_cols > 0:
        col_width_mm = available_width_mm / num_cols
        for col in t.columns:
            col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    doc.add_paragraph()


def parse_compliance_summary() -> List[dict]:
    """
    Parse the embedded compliance summary CSV for determinands
    with defined regulatory parameters (including pesticides).
    """
    compliance_csv = """Determinand,Reg_Determinand_Name,Reg_PCV_Value,Reg_PCV_Units,Compliance_Pass_Fail,Total_Samples,Num_Fail,Mean_Result,Max_Result
1 2-Dichloroethane (Total),,,,,,,,
Aluminium (Total),Aluminium,200.0,µg/l,Fail,4464.0,2.0,9.45493,251.5
Ammonium (Total),Ammonium,0.5,mgNH4/l,Pass,2921.0,0.0,0.155883,0.482
Antimony,Antimony,5.0,µg/l,Pass,1322.0,0.0,0.180499,0.8
Arsenic (Total),,,,,,,,
Benzene (Total),,,,,,,,
Benzo[a]Pyrene (Total),benzo[a]pyrene,0.01,µg/l,Pass,1295.0,0.0,0.00201776,0.008
Boron,Boron,1.0,mg/l,Pass,49.0,0.0,0.0612796,0.187
Bromate,Bromate,10.0,µg/l,Pass,142.0,0.0,0.759366,2.04
Cadmium (Total),,,,,,,,
Chloride,Chloride,250.0,mg/l,Pass,51.0,0.0,48.3482,78.9
Chromium (Total),Chromium,50.0,µg/l,Pass,1321.0,0.0,0.550683,4.0
Clostridum Perfringens (Sulphite-reducing Clostridia) (Confirmed),Clostridium perfringens,0.0,Number/100ml,Pass,51.0,0.0,0.0,0.0
Coliform Bacteria (Indicator),Coliform bacteria,0.0,number/100ml,Fail,13375.0,41.0,0.039028,95.0
Colony Counts After 3 Days At 22øc (Colony Counts),,,,,,,,
Colour,Colour,20.0,mg/l Pt/Co,Pass,5057.0,0.0,1.30739,4.4
Conductivity (Electrical Conductivity),Conductivity,2500.0,µS/cm @ 20°C,Pass,1920.0,0.0,649.43,887.0
Copper (Total),Copper,2.0,mg/l,Pass,1279.0,0.0,0.0684758,1.23
Cyanide (Total),Cyanide,50.0,µg/l,Pass,52.0,0.0,3.333,3.333
E.Coli (faecal coliforms Confirmed),Escherichia coli,0.0,number/100ml,Fail,13375.0,4.0,0.00747664,89.0
Enterococci (Confirmed),Enterococci,0.0,number/100ml,Fail,1314.0,2.0,0.00152207,1.0
Fluoride (Total),Fluoride,1.5,mg/l,Pass,1329.0,0.0,0.381083,1.366
Gross Alpha,gross alpha,0.1,Bq/l,Fail,9.0,1.0,0.0552222,0.152
Gross Beta,gross beta,1.0,Bq/l,Pass,9.0,0.0,0.185889,0.3
Hydrogen ion (pH) - Indicator (Hydrogen ion) (pH),pH,,,,,,,
Iron (Total),Iron,200.0,µg/l,Fail,4484.0,7.0,17.0205,1495.0
Lead (10 - will apply 25.12.2013),Lead,10.0,µg/l,Fail,1325.0,5.0,0.463983,39.669
Manganese (Total),Manganese,50.0,µg/l,Fail,4484.0,2.0,3.12674,185.9
Mercury (Total),Mercury,1.0,µg/l,Pass,51.0,0.0,0.0920588,0.097
Nickel (Total),Nickel,20.0,µg/l,Fail,1320.0,6.0,2.81667,299.83
Nitrate (Total),Nitrate,50.0,mgNO3/l,Pass,2928.0,0.0,19.5823,49.0
Nitrite - Consumer's Taps,Nitrite,0.5,mgNO2/l,Pass,2937.0,0.0,0.0619294,0.407
Nitrite/Nitrate formula,Nitrate+Nitrite formula,1.0,,Pass,2863.0,0.0,0.381151,0.98
Odour,Odour,,,,,,,
Pesticides (Total by Calculation),Pesticides: total,0.5,µg/l,Pass,45.0,0.0,0.00872222,0.037
Pesticides 2 4-D (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.013,0.016
Pesticides 2 4-DB (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.008,0.008
Pesticides Aldrin,Aldrin,0.03,µg/l,Pass,30.0,0.0,0.00296,0.003
Pesticides Atrazine (Total),Other pesticides,0.1,µg/l,Pass,11.0,0.0,0.004,0.004
Pesticides Benazolin (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.01,0.01
Pesticides Bentazone (Total),Other pesticides,0.1,µg/l,Pass,7.0,0.0,0.008,0.008
Pesticides Bromacil,Other pesticides,0.1,µg/l,Pass,5.0,0.0,0.006,0.006
Pesticides Carbetamide,Other pesticides,0.1,µg/l,Pass,5.0,0.0,0.014,0.014
Pesticides Chlorfenvinphos (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.007,0.007
Pesticides Chloridazon,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.012,0.012
Pesticides Chlormequat,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.005,0.005
Pesticides Chlorothalonil,Other pesticides,0.1,µg/l,Pass,1.0,0.0,0.01,0.01
Pesticides Chlortoluron (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.01,0.01
Pesticides Clopyralid (Total),Other pesticides,0.1,µg/l,Pass,14.0,0.0,0.018,0.018
Pesticides Cyanazine,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.008,0.008
Pesticides Dichlorprop (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.004,0.004
Pesticides Dieldrin (Total),Dieldrin,0.03,µg/l,Pass,29.0,0.0,0.00296897,0.003
Pesticides Fluroxypyr,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.02,0.02
Pesticides Glyphosate,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.007,0.007
Pesticides Heptachlor (Total),Heptachlor,0.1,µg/l,Pass,30.0,0.0,0.002,0.002
"Pesticides Heptachlor Epoxide - Total (Trans, CIS) (Heptachlor Epoxide)",Heptachlor epoxide,0.1,µg/l,Pass,30.0,0.0,0.003,0.003
Pesticides Isoproturon (Total),Other pesticides,0.1,µg/l,Pass,5.0,0.0,0.006,0.006
Pesticides Linuron (Total),Other pesticides,0.1,µg/l,Pass,24.0,0.0,0.014,0.014
Pesticides MCPA (Total) 4-chloro-o-tolyloxyacetic acid,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.008,0.008
Pesticides MCPB (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.006,0.006
Pesticides MCPP(Mecoprop) (Total),Other pesticides,0.1,µg/l,Pass,7.0,0.0,0.006,0.006
Pesticides Metaldehyde,Other pesticides,0.1,µg/l,Pass,18.0,0.0,0.0208167,0.037
Pesticides Metamitron,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.017,0.017
Pesticides Metazachlor,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.03,0.03
Pesticides Monuron,Other pesticides,0.1,µg/l,Pass,24.0,0.0,0.012,0.012
Pesticides Pirimicarb,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.00665,0.007
Pesticides Prometryne,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.01,0.01
Pesticides Propazine,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.01,0.01
Pesticides Propyzamide (Total),Other pesticides,0.1,µg/l,Pass,5.0,0.0,0.0126,0.024
Pesticides Quinmerac,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.013,0.013
Pesticides Simazine (Total),Other pesticides,0.1,µg/l,Pass,5.0,0.0,0.01,0.01
Pesticides Terbutryn,Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.01,0.01
Pesticides Triclopyr (Total),Other pesticides,0.1,µg/l,Pass,2.0,0.0,0.012,0.012
Pesticides Trietazine,Other pesticides,0.1,µg/l,Pass,24.0,0.0,0.016,0.016
Polycyclic Aromatic Hydrocarbons (Total by Calculation),PAH,0.1,µg/l,Pass,1293.0,0.0,2.78422e-05,0.036
Residual Disinfectant - Free,,,,,,,,
Residual Disinfectant - Total,,,,,,,,
Selenium (Total),Selenium,10.0,µg/l,Pass,1322.0,0.0,0.901224,3.68
Sodium (Total),Sodium,200.0,mg/l,Pass,1322.0,0.0,32.2581,98.6
Sulphate,Sulphate,250.0,mg/l,Pass,50.0,0.0,60.186,125.54
Taste (Taste Quant),Taste,,,,,,,
Tetrachloromethane (Total),,,,,,,,
Total Organic Carbon,TOC,,,,,,,
Trichloroethene & Tetrachloroethene - Sum Of 2 Substances (Total by Calculation),,,,,,,,
Trihalomethanes (Total by Calculation),THM (total),100.0,µg/l,Pass,1324.0,0.0,15.9668,80.0
Tritium,tritium,100.0,Bq/l,Pass,9.0,0.0,5.31444,6.0
Turbidity,Turbidity,1.0,NTU,Fail,5058.0,17.0,0.135258,2.73
"""
    f = StringIO(compliance_csv)
    reader = csv.DictReader(f)
    records: List[dict] = []
    for row in reader:
        reg_name = row.get("Reg_Determinand_Name", "").strip()
        if reg_name == "":
            continue
        records.append(row)
    return records


def format_number(value: str, decimals: int = 2, allow_blank: bool = True) -> str:
    if value is None:
        return "" if allow_blank else "0"
    text = str(value).strip()
    if text == "":
        return "" if allow_blank else "0"
    try:
        num = float(text)
    except ValueError:
        return text
    if num.is_integer():
        return f"{int(num):,d}"
    fmt = "{:,.%df}" % decimals
    return fmt.format(num)


def build_compliance_tables(records: List[dict]) -> tuple[list[list[str]], list[list[str]]]:
    """
    Build two 5-column tables from the compliance records
    to meet the column limit while covering all requested fields.
    """
    header1 = [
        "Determinand",
        "Reg. name",
        "PCV value",
        "PCV units",
        "Compliance",
    ]
    header2 = [
        "Determinand",
        "Total samples",
        "Failures",
        "Mean result",
        "Max result",
    ]
    rows1 = [header1]
    rows2 = [header2]

    for row in records:
        det = row.get("Determinand", "").strip()
        reg_name = row.get("Reg_Determinand_Name", "").strip()
        pcv_val = format_number(row.get("Reg_PCV_Value", ""), 3)
        pcv_units = row.get("Reg_PCV_Units", "").strip()
        comp = row.get("Compliance_Pass_Fail", "").strip()

        total_samples = format_number(row.get("Total_Samples", ""), 0)
        num_fail = format_number(row.get("Num_Fail", ""), 0)
        mean_res = format_number(row.get("Mean_Result", ""), 4)
        max_res = format_number(row.get("Max_Result", ""), 4)

        rows1.append([det, reg_name, pcv_val, pcv_units, comp])
        rows2.append([det, total_samples, num_fail, mean_res, max_res])

    return rows1, rows2


def build_failed_summary_table(records: List[dict]) -> list[list[str]]:
    """
    Build a 5-column summary table containing only determinands
    that failed compliance, without the 'Reg. name' column.
    Columns: Determinand, PCV (value and units), Total samples, Failures, Max result.
    """
    header = [
        "Determinand",
        "PCV (value and units)",
        "Total samples",
        "Failures",
        "Max result",
    ]
    rows = [header]

    for row in records:
        if row.get("Compliance_Pass_Fail", "").strip().lower() != "fail":
            continue

        det = row.get("Determinand", "").strip()
        pcv_val = format_number(row.get("Reg_PCV_Value", ""), 3)
        pcv_units = row.get("Reg_PCV_Units", "").strip()
        if pcv_units:
            pcv_combined = f"{pcv_val} {pcv_units}"
        else:
            pcv_combined = pcv_val

        total_samples = format_number(row.get("Total_Samples", ""), 0)
        num_fail = format_number(row.get("Num_Fail", ""), 0)
        max_res = format_number(row.get("Max_Result", ""), 4)

        rows.append([det, pcv_combined, total_samples, num_fail, max_res])

    return rows


def generate_report() -> None:
    doc = Document()
    configure_page(doc)

    report_title = "Anglian Water 2022 Domestic Water Quality – UK Regulatory Compliance Assessment"
    today_str = date.today().strftime("%d %B %Y")

    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(doc, include_toc=True, include_lof=True, include_lot=True)

    # 1 Introduction
    doc.add_heading("1 Introduction", level=1)
    para = doc.add_paragraph()
    para.style = "Body Text"
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.add_run(
        "This report evaluates Anglian Water’s 2022 domestic potable water quality "
        "results against the requirements of the Water Supply (Water Quality) "
        "Regulations in England. The assessment uses the Anglian Water domestic "
        "water quality monitoring dataset for consumer taps, focusing on compliance "
        "with prescribed concentration or value (PCV) limits for microbiological, "
        "chemical, radiological and aesthetic parameters, including all regulated "
        "pesticides."
    )

    # 2 Objectives
    doc.add_heading("2 Objectives", level=1)
    para = doc.add_paragraph()
    para.style = "Body Text"
    para.add_run(
        "The specific objectives of this compliance assessment are set out below."
    )
    obj_items = [
        "Check compliance of all determinands with applicable UK Water Supply (Water Quality) Regulations PCVs for 2022 domestic tap samples.",
        "Summarise, for each determinand, the regulatory requirement, observed mean and maximum concentrations, total number of samples, number of failures and overall pass/fail status, including all pesticide parameters.",
        "Characterise temporal behaviour of any non-compliant determinands using time series plots with PCV limits where relevant.",
        "Identify patterns that may indicate treatment process issues, distribution system impacts or premise-specific plumbing effects.",
        "Provide practical recommendations and suggestions for further analysis, sampling or operational investigation."
    ]
    for item in obj_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # 3 Description of Dataset
    doc.add_heading("3 Description of Dataset", level=1)
    para = doc.add_paragraph(style="Body Text")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.add_run(
        "The assessment is based on the file "
        "\"Anglian_Water_Domestic_Water_Quality.csv\", containing 130,664 individual "
        "laboratory results from consumer tap samples taken between 03/01/2022 and "
        "30/12/2022. Each record is identified by a unique Sample_Id and includes "
        "sampling date and time, DWI determinand code, determinand description, "
        "unit, operator qualifier, numeric result, Lower Super Output Area (LSOA) "
        "code and provider metadata. Sampling is irregular, with multiple "
        "determinands reported per sample event and frequencies varying by parameter."
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Regulatory compliance in this report uses a pre-processed summary "
        "table that maps determinands to their statutory PCVs and aggregates "
        "sample counts, failures and summary statistics."
    )

    # 4 Analysis
    doc.add_heading("4 Analysis", level=1)

    # 4.1 Overall compliance against UK PCVs
    doc.add_heading("4.1 Overall compliance against UK PCVs", level=2)
    para = doc.add_paragraph(style="Body Text")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.add_run(
        "Table 1 summarises only those determinands that recorded at least one "
        "sample above their prescribed concentration or value (PCV) in 2022. "
        "The table highlights the regulatory requirement, monitoring intensity and "
        "extent of exceedance for each non-compliant parameter. The full compliance "
        "suite for all determinands, including pesticides and parameters with zero "
        "failures, is provided in Tables 2 and 3 in Appendix A."
    )

    records = parse_compliance_summary()
    failed_summary_rows = build_failed_summary_table(records)
    add_table(
        doc,
        "Summary of determinands with PCV failures.",
        failed_summary_rows,
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "The summary shows that non-compliance is restricted to a small group of "
        "determinands: microbiological indicators (coliform bacteria, Escherichia "
        "coli and enterococci), turbidity, certain aesthetic or operational metals "
        "(aluminium, iron and manganese), lead and nickel at consumer taps, and a "
        "single gross alpha screening exceedance. All other determinands listed in "
        "Appendix A comply fully with their statutory limits."
    )

    # 4.2 Microbiological determinands
    doc.add_heading("4.2 Microbiological determinands", level=2)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Microbiological performance is a primary public health concern. For coliform "
        "bacteria, E. coli and enterococci the PCV is 0 number/100 ml, so any "
        "detection constitutes non-compliance at the sample level. The dataset "
        "includes 13,375 samples for coliform bacteria and E. coli and 1,314 samples "
        "for enterococci."
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Coliform bacteria failures (41 out of 13,375 samples) are scattered "
        "throughout the year, with a small number of pronounced spikes indicating "
        "localised contamination or sampling events rather than sustained problems."
    )
    try:
        doc.add_picture("coliform_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Coliform bacteria time series at consumer taps (PCV 0 number/100 ml).",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Spatial patterns of coliform exceedance are illustrated in Figure 2, which "
        "maps failure rates by LSOA. Almost all LSOAs show very low exceedance "
        "fractions, typically well below 0.1, with a few small clusters approaching "
        "0.2–0.4 where monitoring density is often low. This supports the view that "
        "microbiological non-compliance is restricted to isolated locations."
    )
    try:
        doc.add_picture("coliform_exceedance_rate_by_lsoa_map.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Coliform bacteria exceedance rate by LSOA, 2022 (capped at 0.5 failures/total samples).",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "E. coli results are overwhelmingly compliant: only four failures are "
        "observed out of 13,375 samples. As shown in Figure 3, three discrete "
        "events are visible, including one high outlier around late March "
        "(approximately 90 number/100 ml). These incidents warrant incident-level "
        "review but do not indicate systemic treatment breakdown."
    )
    try:
        doc.add_picture("ecoli_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "E. coli time series at consumer taps (PCV 0 number/100 ml).",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Enterococci monitoring shows a single low-level failure (1 number/100 ml) "
        "within 1,314 samples. Figure 4 confirms that this was an isolated event, "
        "consistent with a transient contamination or sampling artefact rather than "
        "a persistent faecal contamination issue."
    )
    try:
        doc.add_picture("enterococci_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Enterococci time series at consumer taps (PCV 0 number/100 ml).",
        )
    except Exception:
        pass

    # 4.3 Turbidity and aesthetic/operational metals
    doc.add_heading("4.3 Turbidity and aesthetic/operational metals", level=2)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Operational parameters such as turbidity and metals (aluminium, iron and "
        "manganese) are sensitive indicators of treatment performance and "
        "distribution system condition. While exceedances mainly affect appearance "
        "and customer acceptability, high values can correlate with particle "
        "breakthrough and microbiological risk if not well controlled."
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Turbidity exhibits generally low values, with most results between about "
        "0.05 and 0.4 NTU. The PCV is 1 NTU. Figure 5 shows infrequent exceedances "
        "above this limit (17 out of 5,058 samples) and a small number of higher "
        "outliers up to approximately 2.7 NTU. These peaks are scattered through "
        "the year, suggesting short-lived local disturbances such as mains "
        "operations, flushing or filter backwash carry-over."
    )
    try:
        doc.add_picture("turbidity_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Turbidity time series with UK PCV limit of 1 NTU.",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Aluminium performance is generally good, but the compliance summary "
        "records two failures out of 4,464 samples against the 200 µg/l PCV. "
        "Figure 6 reveals a distinct mid-year period (approximately March to July) "
        "with elevated aluminium concentrations and several results approaching or "
        "exceeding the PCV, including one value slightly above 250 µg/l. After "
        "August, levels return to a low baseline, implying that process conditions "
        "at one or more treatment works were temporarily sub-optimal during that "
        "period."
    )
    try:
        doc.add_picture("aluminium_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Aluminium time series with PCV 200 µg/l.",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Iron results are predominantly below 80 µg/l, yet seven samples out of "
        "4,484 exceed the 200 µg/l PCV. Figure 7 shows sporadic high events during "
        "May–June, August and late in the year. One extreme outlier exceeds "
        "1,400 µg/l, indicative of significant deposit mobilisation or disturbance "
        "in local iron pipework or storage assets. The absence of sustained high "
        "background levels suggests that these are distribution-side rather than "
        "treatment works problems."
    )
    try:
        doc.add_picture("iron_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Iron time series with PCV 200 µg/l.",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Manganese concentrations are mostly between about 0.5 and 6 µg/l, well "
        "below the 50 µg/l PCV. Only two failures occur among 4,484 samples. "
        "Figure 8 highlights a small number of pronounced spikes, including a peak "
        "around 185 µg/l in late June–July and another high result near 125 µg/l. "
        "These are consistent with localised deposit mobilisation in low-turnover "
        "mains or reservoirs."
    )
    try:
        doc.add_picture("manganese_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Manganese time series with PCV 50 µg/l.",
        )
    except Exception:
        pass

    # 4.4 Consumer tap metals: lead and nickel
    doc.add_heading("4.4 Consumer tap metals: lead and nickel", level=2)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Lead and nickel are primarily associated with premise and distribution "
        "pipework materials rather than treatment processes. Compliance performance "
        "is therefore strongly influenced by plumbing configuration, stagnation and "
        "asset replacement programmes."
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Lead exhibits very low concentrations for the majority of samples "
        "(typically below 2 µg/l). Five failures are recorded out of 1,325 samples "
        "against the 10 µg/l PCV. Figure 9 shows isolated exceedances spread across "
        "the year, including one high outlier approaching 40 µg/l in late autumn. "
        "The pattern indicates premise-specific plumbing or occasional stagnation "
        "rather than a systemic issue."
    )
    try:
        doc.add_picture("lead_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Lead time series at consumer taps with PCV 10 µg/l.",
        )
    except Exception:
        pass

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Nickel performance is similarly strong overall, with most results below "
        "5–6 µg/l and only six failures among 1,320 samples when compared with the "
        "20 µg/l PCV. Figure 10 shows a few moderate exceedances in late summer to "
        "autumn and a single very high value approaching 300 µg/l. These "
        "concentrations are consistent with localised nickel-containing fittings "
        "or stagnation in particular properties."
    )
    try:
        doc.add_picture("nickel_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Nickel time series at consumer taps with PCV 20 µg/l.",
        )
    except Exception:
        pass

    # 4.5 Radiological parameters
    doc.add_heading("4.5 Radiological parameters", level=2)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Radiological screening is represented by gross alpha and gross beta "
        "activity measurements, supported by tritium monitoring. The regulatory "
        "approach uses screening levels (0.1 Bq/l for gross alpha and 1.0 Bq/l for "
        "gross beta) to determine whether more detailed radionuclide-specific "
        "assessments are required."
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "All gross beta and tritium results are comfortably below their screening "
        "values. For gross alpha, eight of nine samples lie between about 0.02 and "
        "0.053 Bq/l, but one sample at LSOA E01017616 on 25 January 2022 reaches "
        "0.152 Bq/l, exceeding the 0.1 Bq/l screening threshold. This is treated as "
        "a failure in the summary table, triggering follow-up investigation as "
        "required by the Regulations."
    )
    try:
        doc.add_picture("grossalpha_timeseries.png", width=Mm(160))
        add_caption(
            doc,
            "Figure",
            "Gross alpha time series with screening value 0.1 Bq/l.",
        )
    except Exception:
        pass

    # 4.6 Pesticides and organic contaminants
    doc.add_heading("4.6 Pesticides and organic contaminants", level=2)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "The dataset includes a comprehensive suite of regulated pesticides and "
        "related organic contaminants. As detailed in the full compliance tables in "
        "Appendix A, all pesticide determinands, the total pesticides parameter and "
        "trihalomethanes comply with their respective PCVs. Observed concentrations "
        "are typically one to two orders of magnitude below the 0.1 µg/l individual "
        "pesticide limit and the 0.5 µg/l total pesticides limit, with low "
        "variability across the monitoring period."
    )

    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "Polycyclic aromatic hydrocarbons, benzo[a]pyrene and related organic "
        "indicators are also consistently below their regulatory limits, implying "
        "effective control of raw water quality, treatment processes and material "
        "selection within the distribution system."
    )

    # 5 Key Findings
    doc.add_heading("5 Key Findings", level=1)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "The principal findings from the 2022 Anglian Water domestic compliance "
        "assessment are summarised below."
    )

    findings = [
        "Overall regulatory compliance is high: most determinands, including all monitored pesticides, nitrate, nitrite, bromate, trihalomethanes and the majority of metals, show zero failures against their statutory PCVs.",
        "Microbiological performance is strong, with very low failure rates for coliform bacteria (41/13,375), E. coli (4/13,375) and enterococci (2/1,314). Failures are sporadic in time and geographically localised, consistent with isolated premise or small-zone events rather than widespread contamination.",
        "Turbidity, aluminium, iron and manganese exhibit occasional exceedances and outliers. Aluminium shows a mid-year cluster of elevated values consistent with a temporary coagulation or filtration issue, while iron and manganese spikes point towards local deposit mobilisation in the distribution system.",
        "Lead and nickel concentrations are generally very low, but isolated exceedances and rare high outliers indicate ongoing exposure risks for a small number of properties, likely associated with legacy plumbing or fittings.",
        "A single gross alpha sample exceeds the 0.1 Bq/l screening value, requiring radionuclide-specific follow-up, but subsequent samples remain below the threshold, suggesting a transient or localised cause rather than a systemic radiological issue.",
        "Spatial mapping of coliform failures demonstrates that LSOAs with the highest exceedance rates tend to have relatively few samples, emphasising the need to interpret apparent high rates in the context of sample numbers and to prioritise zones with both repeated detections and adequate monitoring coverage."
    ]
    for item in findings:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # 6 Recommendations
    doc.add_heading("6 Recommendations", level=1)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "On the basis of the observed patterns and regulatory performance, the "
        "following actions and further analyses are recommended."
    )

    recs = [
        "Investigate individual microbiological failures (coliforms, E. coli and enterococci) using Sample_Id and LSOA to review local plumbing, sampling conditions and any concurrent operational events such as bursts, repairs or changes in supply configuration.",
        "Review treatment works performance during the mid-2022 period of elevated aluminium, including coagulation control, sludge management and filter optimisation, and confirm whether any process changes or corrective actions coincided with the return to baseline levels.",
        "Use event logs, customer contacts and asset records to correlate high iron and manganese results with specific mains, service reservoirs or premise connections, and target cleaning, flushing or asset renewal where repeated spikes occur.",
        "Undertake property-level investigations for addresses associated with lead and nickel exceedances, including plumbing surveys, stagnation assessments and consideration of point-of-use flushing advice or replacement of lead or nickel-containing fittings.",
        "Carry out a radionuclide-specific investigation for the gross alpha exceedance sample, including follow-up sampling and analysis for individual alpha-emitting radionuclides, and assess the resulting indicative dose in line with regulatory guidance.",
        "Extend spatial analysis beyond coliforms to include turbidity, iron, manganese, lead and nickel, using exceedance maps and control charts by supply zone or district metered area to identify recurring hotspots.",
        "Consider targeted additional sampling in LSOAs with high-coliform exceedance rates based on few samples, to distinguish between random single events and genuine local vulnerability.",
        "Develop integrated dashboards linking laboratory data with operational telemetry, works performance indicators and network events, enabling near real-time detection and response to emerging compliance risks."
    ]
    for item in recs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # 7 Appendix A – Detailed compliance tables
    doc.add_heading("7 Appendix A – Detailed compliance tables", level=1)
    para = doc.add_paragraph(style="Body Text")
    para.add_run(
        "This appendix presents the full determinand-by-determinand compliance "
        "summary used in the assessment. Tables 2 and 3 list the regulatory "
        "parameters, monitoring statistics and result distributions for all "
        "determinands with defined PCVs, including those that were fully compliant "
        "throughout 2022."
    )

    table1_rows, table2_rows = build_compliance_tables(records)
    add_table(
        doc,
        "Regulatory limits and compliance status by determinand (part 1).",
        table1_rows,
    )
    add_table(
        doc,
        "Monitoring statistics and result distribution by determinand (part 2).",
        table2_rows,
    )

    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()