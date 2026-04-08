# filename: report_template.py
"""
Template for generating a Word report using python-docx.

IMPORTANT IMPLEMENTATION NOTES (for LLMs):

- ALWAYS create a new document with `Document()` (with parentheses, no arguments).
  - `Document("path")` is ONLY for opening an existing .docx file.
- DO NOT remove required positional arguments from python-docx methods:
  - `Document.add_table(rows, cols, style=None)` → rows and cols are REQUIRED.
  - `_Header.add_table(rows, cols, width)` / `_Footer.add_table(rows, cols, width)`
    → rows and cols are REQUIRED.
- When a function parameter type is a Length (Mm, Pt, Inches, etc.), you MUST pass a Length
  object (e.g. `Pt(10)`, `Mm(25)`), NOT a bare integer or float.
- When setting colors using RGBColor, you MUST use `.font.color.rgb = RGBColor(...)`
  or `.font.color.theme_color = ...`; do NOT assign an RGBColor directly to `font.color`.
- Style names like "Title", "Subtitle", "Caption", "TOC Heading",
  "Medium Shading 1 Accent 1" must match styles that actually exist in the target template.
- Images MUST be inserted from **local file paths only**:
    doc.add_picture("image.png", width=Inches(4))
  Do NOT use file-like objects, BytesIO, URLs or other streams for images.
"""

from datetime import date
from typing import List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# LLM NOTE:
# - WD_ALIGN_VERTICAL is effectively an alias for WD_CELL_VERTICAL_ALIGNMENT.
#   It exposes the same members (TOP, CENTER, BOTTOM, etc.) for cell vertical alignment.


# ────────────────────────────────────────────────────────────────
#                       low-level helpers
# ────────────────────────────────────────────────────────────────
def configure_page(doc: Document) -> None:
    """
    Apply A4 size and 25 mm margins to every section.

    LLM NOTE:
    - `doc.sections` is a read-only collection of Section objects.
    - You must iterate over it (as done below) to configure ALL sections.
    - `Mm(...)` is a Length constructor
      which is the CORRECT type for margin and page size properties.
    """
    for s in doc.sections:
        s.page_height = Mm(297)  # A4 height in mm
        s.page_width = Mm(210)  # A4 width in mm
        s.top_margin = Mm(25)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(25)


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    font_size: Pt | None = None,
    font_color: RGBColor | None = None,
) -> None:
    """
    Add a Word field inside a run so we can style it.

    LLM NOTES:
    - `paragraph` MUST be a python-docx Paragraph instance (e.g. from `doc.add_paragraph()`).
    - `field_code` is the raw Word field instruction string (e.g. 'PAGE  \\* MERGEFORMAT').
    - `font_size` MUST be a Length (e.g. Pt(10)) or None.
      DO NOT pass an int directly (e.g. do NOT use `font_size=10`; use `font_size=Pt(10)`).
    - `font_color` MUST be an RGBColor instance or None.
    - OxmlElement and qn are used here because python-docx has no high-level API
      for arbitrary Word fields; DO NOT remove this low-level XML code.
    """
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run._r.append(fld)

    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        # IMPORTANT: use .rgb when assigning RGBColor
        run.font.color.rgb = font_color


def add_caption(doc: Document, seq_type: str, caption_text: str) -> None:
    """
    Automatic caption beginning with “Figure/Table n – …”.

    *seq_type* must be 'Figure' or 'Table' (case-insensitive).

    LLM NOTES:
    - This function produces captions such as:
        Figure 1 – Caption text
        Figure 2 – Another caption
      or:
        Table 1 – Caption text
    - `seq_type` is used both in the visible text and in the SEQ field.
      It MUST match the label used by List of Figures/Tables fields
      (e.g. "Figure", "Table") for automatic lists to work.
    """
    seq_type = seq_type.title()
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{seq_type} ")
    insert_field(p, f"SEQ {seq_type} \\* ARABIC")
    p.add_run(f" – {caption_text}")


def add_header_footer(
    doc: Document, title: str, report_date: str | None = None
) -> None:
    """
    Add a simple header and a footer with date (left) and page number (right).

    LLM NOTES:
    - This modifies EACH section's header and footer.
    - This uses `_Header.add_table()` / `_Footer.add_table()` which have the
      signature `add_table(rows, cols, width)` and DO accept a `width` Length.
    - Do NOT replace these calls with `Document.add_table(...)`, because
      `Document.add_table` has the signature `add_table(rows, cols, style=None)`
      and does NOT accept a `width` argument.
    - Page number is inserted as a Word field via `insert_field`, but its final
      appearance is controlled by the footer paragraph style (see below).
    """
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    grey = RGBColor(102, 102, 102)
    fsz10 = Pt(10)

    # --- IMPORTANT: style the Footer paragraph style itself ---
    # This ensures that field results (PAGE) also appear as 10pt grey.
    try:
        footer_style = doc.styles["Footer"]
        footer_style.font.size = fsz10
        footer_style.font.color.rgb = grey
    except KeyError:
        # If the "Footer" style does not exist (unusual for a standard template),
        # we silently skip. In that case, footer text may fall back to defaults
        # (typically 11pt black).
        footer_style = None

    for sec in doc.sections:
        # header
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_run = hp.add_run(title)
        hdr_run.font.size = fsz10
        hdr_run.font.color.rgb = grey

        # footer
        tbl = sec.footer.add_table(rows=1, cols=2, width=Mm(160))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in tbl.rows[0].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # date (left cell)
        left_para = tbl.cell(0, 0).paragraphs[0]
        if footer_style is not None:
            left_para.style = footer_style

        d_run = left_para.add_run(report_date)
        d_run.font.size = fsz10
        d_run.font.color.rgb = grey

        # page number (right cell)
        p_par = tbl.cell(0, 1).paragraphs[0]
        if footer_style is not None:
            p_par.style = footer_style

        p_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        insert_field(p_par, "PAGE  \\* MERGEFORMAT", font_size=fsz10, font_color=grey)


def add_title_page(doc: Document, title: str, report_date: str | None = None) -> None:
    """Create a standalone title page followed by a page break."""
    if report_date is None:
        report_date = date.today().strftime("%d %B %Y")

    doc.add_paragraph(title, style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        report_date, style="Subtitle"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()  # IMPORTANT: uses Document.add_page_break(), no arguments.


def add_front_matter(
    doc: Document,
    include_toc: bool = True,
    include_lof: bool = True,
    include_lot: bool = True,
) -> None:
    """
    Insert TOC / List-of-Figures / List-of-Tables – each optional.

    LLM NOTES:
    - Each list is implemented as a TOC field with different switches:
        - Main TOC: 'TOC \\o "1-3" \\h \\z \\u'
        - Figures:   'TOC \\h \\z \\c "Figure"'
        - Tables:    'TOC \\h \\z \\c "Table"'
    - The "Figure" / "Table" labels MUST match the captions generated by add_caption.
    """
    if include_toc:
        doc.add_paragraph("Table of Contents", style="TOC Heading")
        toc = doc.add_paragraph()
        insert_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_paragraph()  # Spacer

    if include_lof:
        doc.add_paragraph("Figures", style="TOC Heading")
        lof = doc.add_paragraph()
        insert_field(lof, 'TOC \\h \\z \\c "Figure"')
        doc.add_paragraph()  # Spacer

    if include_lot:
        doc.add_paragraph("Tables", style="TOC Heading")
        lot = doc.add_paragraph()
        insert_field(lot, 'TOC \\h \\z \\c "Table"')
        doc.add_paragraph()  # Spacer

    if any((include_toc, include_lof, include_lot)):
        doc.add_page_break()


def add_table(
    doc: Document,
    caption: str,
    rows: List[List[str]],
    style: str = "Medium Shading 1 Accent 1",
) -> None:
    """
    Create a formatted table and auto-caption it.

    LLM NOTES:
    - `rows` MUST be a non-empty list of row lists; each row list MUST have the same length.
      Example:
          rows = [
              ["Heading 1", "Heading 2"],
              ["Value 1", "Value 2"],
          ]
      i.e. rows[0] is the header row.
    - Document.add_table requires BOTH `rows` and `cols` positional arguments:
        doc.add_table(rows=len(rows), cols=len(rows[0]))
      DO NOT omit these.
    """
    if not rows:
        return
    add_caption(doc, "Table", caption)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Adjust column widths to fit content and page width.
    # This is a basic attempt; for complex tables, more advanced logic might be needed.
    # A4 width is 210mm, with 25mm margins, leaving ~160mm for content.
    # Distribute available width equally among columns as a starting point.
    available_width_mm = 160
    num_cols = len(rows[0])
    if num_cols > 0:
        col_width_mm = available_width_mm / num_cols
        for col in t.columns:
            # IMPORTANT: width MUST be a Length object; use Mm(), not an int.
            # Correct:  col.width = Mm(col_width_mm)
            # WRONG:    col.width = 20  # bare int, not a Length
            col.width = Mm(col_width_mm)

    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            t.cell(r, c).text = str(txt)

    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer paragraph after the caption


# ────────────────────────────────────────────────────────────────
#                       MAIN skeleton
# ────────────────────────────────────────────────────────────────
def generate_report() -> None:
    """
    Main entry point to generate the report.

    LLM NOTES:
    - ALWAYS create the document with `Document()` (no arguments) to start from the default.
    - ALWAYS call `doc.save("report.docx")` once all content has been added.
    - You may customise content between the "CONTENT AREA" markers below, but should
      reuse helper functions (add_table, add_caption, doc.add_picture, etc.) rather than
      duplicating low-level logic.
    """
    doc = Document()
    configure_page(doc)

    report_title = "YOUR REPORT TITLE"  # 👉 replace or derive dynamically
    today_str = date.today().strftime("%d %B %Y")

    # Cover & prelims
    add_title_page(doc, report_title, today_str)
    add_header_footer(doc, report_title, today_str)
    add_front_matter(
        doc,
        include_toc=True,  # toggle independently as required
        include_lof=True,
        include_lot=True,
    )

    # ───────────────────────────────────────────────────────────
    # ✏️  CONTENT AREA – LLMs MAY MODIFY CODE **ONLY BELOW THIS LINE**
    #     DO NOT change:
    #       - imports at the top of the file
    #       - helper function names, parameters, or return types
    #       - calls to configure_page, add_title_page, add_header_footer, add_front_matter
    # ───────────────────────────────────────────────────────────

    # Example heading
    # doc.add_heading("1 Introduction", level=1)

    # Example normal paragraph
    # p = doc.add_paragraph("This is an example paragraph of body text.")
    # p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Example bullet list item
    # doc.add_paragraph("First bullet item", style="List Bullet")

    # Example figure
    # LLM NOTE: Images MUST be inserted from local file paths only.
    # width/height MUST be Length objects (e.g., Inches(4), Mm(80)).
    #
    # from docx.shared import Inches
    # doc.add_picture("figure1.png", width=Inches(4))
    # add_caption(doc, "Figure", "Example figure caption.")

    # Example table
    # table_rows = [
    #     ["Parameter", "Value"],
    #     ["Flow rate", "10 ML/d"],
    #     ["pH", "7.0"],
    # ]
    # add_table(doc, "Example parameter table.", table_rows)

    # ───────────────────────────────────────────────────────────

    # FINAL STEP: Save the document to disk.
    doc.save("report.docx")


if __name__ == "__main__":
    generate_report()
